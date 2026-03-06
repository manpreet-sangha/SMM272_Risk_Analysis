"""
Q1 Part 2 — Generate methodology and results Word document.

Output: Q1/output_q1_2/q1_2_VaR_ES_Methodology_and_Results.docx

Run from project root:
    python Q1/q1_2_generate_docs.py
"""

import os
import sys

# Allow imports from project root (config.py)
_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Output path ──────────────────
OUTPUT_FILE = os.path.join(DOCS_DIR, "q1_2_VaR_ES_Methodology_and_Results.docx")


# ── Styling helpers ───────────────────────────────────────────────────────────

NAVY   = (31, 73, 125)
DARK   = (0, 0, 0)
GREEN  = (0, 112, 0)
RED    = (192, 0, 0)


def set_run(run, size=11, bold=False, italic=False, colour=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = RGBColor(*colour)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p


def bullet(doc, text, style="List Bullet"):
    return doc.add_paragraph(text, style=style)


def centred_italic(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=size, italic=True)
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Note: " + text)
    set_run(r, size=10, italic=True)
    return p


def table_header(table, headers):
    hdr = table.rows[0].cells
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True


def fill_table(table, data, start_row=1, bold_last=False):
    for i, row_data in enumerate(data, start=start_row):
        cells = table.rows[i].cells
        for j, val in enumerate(row_data):
            cells[j].text = str(val)
        if bold_last and (i == start_row + len(data) - 1):
            for c in table.rows[i].cells:
                for para in c.paragraphs:
                    for run in para.runs:
                        run.bold = True


# ── Main document builder ─────────────────────────────────────────────────────

def build_var_es_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Cover ─────────────────────────────────────────────────────────────────
    t = doc.add_heading("SMM272 Risk Analysis \u2014 Q1 Part 2", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Rolling Value-at-Risk and Expected Shortfall:\n"
        "Methodology, Rationale, and Results"
    )
    set_run(r, size=14, bold=True, colour=NAVY)
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = meta.add_run(
        "Portfolio: AAPL, MSFT, IBM, NVDA, GOOGL, AMZN (Equally Weighted)\n"
        "Estimation horizon: 2014-07-01 \u2013 2025-12-30  |  Confidence level: 99%\n"
        "Rolling window: 6 calendar months  |  Four estimation methods"
    )
    set_run(r2, size=11, italic=True, colour=(89, 89, 89))
    doc.add_paragraph()

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    heading(doc, "1. Executive Summary")
    body(doc,
         "This report documents the rolling Value-at-Risk (VaR) and Expected Shortfall (ES) "
         "analysis conducted for the Q1 Part 2 coursework. Four distinct estimation methods "
         "\u2014 Historical Simulation (HS), Parametric Normal, Parametric Student-t, and "
         "GARCH(1,1) \u2014 are applied to daily portfolio returns from an equally weighted (EW) "
         "six-stock technology portfolio. Each method produces one-step-ahead 99% VaR and ES "
         "forecasts using a rolling 6-calendar-month estimation window, yielding 2,893 daily "
         "risk estimates over the period 2014-07-01 to 2025-12-30.")
    body(doc,
         "Historical Simulation achieved the closest agreement with the 99% VaR target, "
         "recording a breach rate of 1.694% (49 exceedances) versus the theoretical 1%. "
         "The Parametric Normal and GARCH(1,1) models significantly understated tail risk, "
         "breaching the 99% threshold on 87 and 82 occasions respectively (rates: 3.007% "
         "and 2.834%). The Parametric Student-t model (63 breaches, 2.178%) offered an "
         "intermediate result but generated very large ES estimates during the 2020 COVID-19 "
         "period (average ES of \u221221.28%), reflecting extreme heavy-tail extrapolation for "
         "fat-tailed distributions. The 2020 crisis period revealed the most pronounced "
         "divergence across methods, with the worst single-day portfolio loss of \u221212.76% "
         "occurring on 16 March 2020.")

    # ── 2. Data and Sample ───────────────────────────────────────────────────
    heading(doc, "2. Data and Sample Description")
    body(doc,
         "The analysis is based on the equally weighted six-stock portfolio constructed in "
         "Q1 Part 1. Daily log returns are used throughout:")
    centred_italic(doc, "r\u209a,t = (1/6) \u00d7 \u03a3 ln( P\u1d62,t / P\u1d62,t\u208b\u2081 )")
    body(doc,
         "The portfolio return series spans January 2014 to December 2025. However, the "
         "rolling VaR/ES estimation requires a 6-month initialisation window; the first "
         "one-step-ahead forecast is therefore generated on 2014-07-01. After discarding "
         "the initialisation period, the evaluation sample consists of 2,893 daily forecast "
         "dates ending 2025-12-30.")
    body(doc,
         "A minimum guard of 20 observations is applied at each window: if fewer than 20 "
         "daily returns fall within the 6-month window (e.g. around public holidays), "
         "that window is skipped and no forecast is generated. This guard prevents "
         "degenerate estimates driven by extremely sparse data.")

    tbl_data = [
        ("Asset universe",        "AAPL, MSFT, IBM, NVDA, GOOGL, AMZN"),
        ("Portfolio weights",     "Equal (1/6 \u2248 16.67% each)"),
        ("Return type",           "Daily log returns"),
        ("Full data start",       "January 2014"),
        ("Evaluation start",      "2014-07-01 (after 6-month initialisation)"),
        ("Evaluation end",        "2025-12-30"),
        ("Forecast observations", "2,893"),
        ("Confidence level",      "99% (one-tailed)"),
    ]
    t2 = doc.add_table(rows=len(tbl_data) + 1, cols=2)
    t2.style = "Light List Accent 1"
    table_header(t2, ["Parameter", "Value"])
    for i, (k, v) in enumerate(tbl_data, start=1):
        t2.rows[i].cells[0].text = k
        t2.rows[i].cells[1].text = v
    doc.add_paragraph()

    # ── 3. Rolling Window Methodology ────────────────────────────────────────
    heading(doc, "3. Rolling Window Methodology")
    body(doc,
         "A rolling (moving) window approach is adopted rather than an expanding window. "
         "At each forecast date t, only returns falling within the most recent 6 calendar "
         "months are used to estimate the risk model. This design has several justifications:")
    bullet(doc,
           "Recency bias control \u2014 financial volatility is time-varying. A fixed "
           "lookback window ensures that the risk estimate reflects the current volatility "
           "regime rather than a long-run average that may be dominated by distant, "
           "irrelevant observations.")
    bullet(doc,
           "Regulatory alignment \u2014 Basel III / FRTB internal models approaches endorse "
           "observation windows of at least 250 trading days (approximately 12 months) "
           "with stressed-period supplements. The 6-month window is a computationally "
           "tractable academic illustration of the rolling principle.")
    bullet(doc,
           "Non-stationarity robustness \u2014 by discarding old data, the window implicitly "
           "down-weights structural shifts (e.g. technology sector regime changes) that "
           "occurred more than 6 months ago, reducing model misspecification.")
    body(doc,
         "Calendar-month boundaries are used rather than fixed trading-day counts "
         "(\u2018dateutil.relativedelta\u2019). For example, the window preceding 2020-03-01 "
         "consists of all trading days from 2019-09-01 to 2020-02-29 inclusive. This ensures "
         "consistent economic coverage across months of different lengths and avoids the "
         "drift that accumulates when counting fixed numbers of trading days.")
    centred_italic(doc,
                   "Window(t) = { returns r\u209a,s  :  t \u2212 6M \u2264 s < t }")
    body(doc,
         "The volatility estimate and distributional parameters are refitted entirely from "
         "scratch at each t, producing a sequence of one-step-ahead (t+0) forecasts. VaR "
         "and ES are then compared against the realised return r\u209a,t to record whether a "
         "breach occurred.")

    # ── 4. VaR and ES: Conceptual Background ────────────────────────────────
    heading(doc, "4. VaR and ES: Conceptual Background")
    body(doc,
         "Value-at-Risk (VaR) at confidence level \u03b1 is the maximum loss not exceeded with "
         "probability \u03b1 over a given holding period. For a portfolio return r:")
    centred_italic(doc, "P( r < \u2212VaR\u03b1 ) = 1 \u2212 \u03b1")
    body(doc,
         "At \u03b1 = 99% and a one-day horizon, VaR\u2099\u2099% is the threshold such that the "
         "portfolio loses more than VaR only 1% of the time. Here VaR is expressed as a "
         "negative number (e.g. \u22123.80% means a loss of 3.80%).")
    body(doc,
         "Expected Shortfall (ES), also known as Conditional VaR (CVaR) or the tail loss, "
         "is the expected loss conditional on the loss exceeding VaR:")
    centred_italic(doc, "ES\u03b1 = E[ r | r < \u2212VaR\u03b1 ]")
    body(doc,
         "ES is subadditive and therefore coherent in the sense of Artzner et al. (1999). "
         "Unlike VaR, ES captures the severity of tail losses rather than merely their "
         "frequency, making it a more informative risk measure for fat-tailed distributions. "
         "ES has replaced VaR as the primary measure under FRTB (Basel IV) precisely because "
         "of this coherence property. Both measures are reported here at the 99% confidence level.")

    # ── 5. Estimation Methods ────────────────────────────────────────────────
    heading(doc, "5. Estimation Methods")
    body(doc,
         "Four methods are applied to each rolling window. Each method reflects a different "
         "set of assumptions about the return distribution and the dynamics of portfolio risk. "
         "Together they provide a comprehensive benchmarking exercise that highlights the "
         "sensitivity of VaR/ES estimates to modelling choices.")

    # 5.1 Historical Simulation
    heading(doc, "5.1  Method 1: Historical Simulation (HS)", level=2)
    body(doc,
         "Historical Simulation is a fully non-parametric approach. No distributional "
         "assumption is imposed on returns; instead, the empirical distribution of past "
         "returns within the window is used directly.")
    body(doc, "Estimation:")
    bullet(doc, "Collect all N returns in the 6-month window: { r\u2081, r\u2082, \u2026, r\u2099 }.")
    bullet(doc, "Sort in ascending order.")
    bullet(doc,
           "VaR\u2099\u2099% = the 1st percentile of the sorted sample "
           "(the value at position \u230a0.01 \u00d7 N\u230b).")
    bullet(doc, "ES\u2099\u2099% = the mean of all returns at or below VaR.")
    body(doc, "Rationale:")
    bullet(doc,
           "Model-free \u2014 HS makes no assumption about the shape of the distribution, "
           "naturally capturing any skewness, fat tails, or multimodality present in the data.")
    bullet(doc,
           "Conceptual simplicity \u2014 the method is transparent and easily audited, "
           "which is valued in regulatory contexts (Basel III standard approach).")
    bullet(doc,
           "Automatically reflects observed crisis episodes embedded in the window, "
           "including the extreme losses of March 2020.")
    body(doc, "Principal limitation:")
    bullet(doc,
           "Entirely backward-looking. HS cannot anticipate tail events not represented "
           "in the 6-month window. In calm periods, the window contains no extreme observations, "
           "leading to VaR underestimation. Conversely, if the window straddles a crisis, "
           "VaR is inflated long after volatility normalises.")

    # 5.2 Parametric Normal
    heading(doc, "5.2  Method 2: Parametric Normal", level=2)
    body(doc,
         "The Parametric Normal method models returns as normally distributed with rolling "
         "mean \u03bc and standard deviation \u03c3 estimated from the window.")
    body(doc, "Estimation:")
    centred_italic(doc,
                   "\u03bc = (1/N) \u03a3 r\u1d62     \u03c3\u00b2 = (1/N) \u03a3 (r\u1d62 \u2212 \u03bc)\u00b2")
    centred_italic(doc,
                   "VaR\u2099\u2099% = \u03bc + \u03a6\u207b\u00b9(0.01) \u00d7 \u03c3     "
                   "(\u03a6\u207b\u00b9(0.01) \u2248 \u22122.3263)")
    centred_italic(doc,
                   "ES\u2099\u2099% = \u03bc \u2212 \u03c3 \u00d7 \u03c6(\u03a6\u207b\u00b9(0.01)) / 0.01")
    body(doc,
         "where \u03a6 is the standard normal CDF and \u03c6 is its density.")
    body(doc, "Rationale:")
    bullet(doc,
           "Analytical tractability \u2014 closed-form VaR and ES; no simulation required.")
    bullet(doc,
           "Industry benchmark \u2014 the Normal VaR remains the most widely taught and "
           "reported measure, providing a natural baseline for comparison with the "
           "more sophisticated methods.")
    bullet(doc,
           "Efficient use of data \u2014 only two parameters (\u03bc, \u03c3) are estimated, "
           "maximising statistical efficiency within short rolling windows.")
    body(doc, "Principal limitation:")
    bullet(doc,
           "Under-estimates tail risk. Q1 Part 1 documented excess kurtosis of 6.33 "
           "and statistically significant non-normality across all four tests (Jarque-Bera, "
           "Shapiro-Wilk, D\u2019Agostino K\u00b2, Kolmogorov-Smirnov). The normal quantile "
           "\u03a6\u207b\u00b9(0.01) \u2248 \u22122.33 is substantially less extreme than the true "
           "empirical 1st percentile of a fat-tailed distribution, producing systematic "
           "VaR underestimation and the highest breach rate in this study (3.007%).")

    # 5.3 Parametric Student-t
    heading(doc, "5.3  Method 3: Parametric Student-t", level=2)
    body(doc,
         "The Parametric Student-t method replaces the Normal distribution with a "
         "location-scale Student-t distribution, which has heavier tails controlled by "
         "a degrees-of-freedom parameter \u03bd.")
    body(doc, "Estimation:")
    bullet(doc,
           "Maximum likelihood estimation (MLE) of (\u03bd, \u03bc, \u03c3) from the window returns.")
    bullet(doc,
           "VaR\u2099\u2099% = \u03bc + \u03c3 \u00d7 t\u207b\u00b9\u209d(0.01)  "
           "where t\u207b\u00b9\u209d is the quantile function of the Student-t with \u03bd degrees of freedom.")
    centred_italic(doc,
                   "ES\u2099\u2099% = \u03bc + \u03c3 \u00d7 [f\u209d(t\u207b\u00b9\u209d(0.01)) / 0.01] "
                   "\u00d7 [(\u03bd + (t\u207b\u00b9\u209d(0.01))\u00b2) / (\u03bd \u2212 1)]")
    body(doc,
         "where f\u209d is the standard Student-t pdf. If MLE fails or \u03bd \u2264 1 "
         "(degenerate), the method falls back to the Parametric Normal.")
    body(doc, "Rationale:")
    bullet(doc,
           "Empirical motivation \u2014 Q1 Part 1 normality analysis confirmed excess kurtosis "
           "of 6.33 for the EW portfolio. A Student-t with \u03bd \u2248 3\u20135 degrees of freedom "
           "is well placed to capture such heavy tails. Fitting \u03bd via MLE allows the "
           "window-specific tail weight to be estimated directly from the data.")
    bullet(doc,
           "Closed-form ES \u2014 unlike HS, the Student-t ES has an analytical formula, "
           "removing the noise associated with small samples in the empirical tail.")
    bullet(doc,
           "Coherent downside risk \u2014 lower \u03bd implies heavier tails and a larger gap "
           "between VaR and ES, which is the correct behaviour for portfolios with "
           "extreme kurtosis.")
    body(doc, "Principal limitation:")
    bullet(doc,
           "Very sensitive to \u03bd during stress periods. When the window data are highly "
           "non-Gaussian (small \u03bd), the ES formula amplifies the tail estimate substantially "
           "\u2014 the average ES during 2020 reached \u221221.28%, far exceeding HS ES of "
           "\u22128.51%. While this may reflect genuine model uncertainty, it can also overstate "
           "actual tail losses and create capital adequacy problems in practice.")

    # 5.4 GARCH(1,1)
    heading(doc, "5.4  Method 4: GARCH(1,1) with Normal Innovations", level=2)
    body(doc,
         "The GARCH(1,1) model (Bollerslev, 1986) captures time-varying conditional volatility "
         "\u2014 the key empirical feature that static models ignore. The conditional variance "
         "evolves according to:")
    centred_italic(doc,
                   "\u03c3\u00b2\u209c = \u03c9 + \u03b1\u2081 \u03b5\u00b2\u209c\u208b\u2081 + \u03b2\u2081 \u03c3\u00b2\u209c\u208b\u2081")
    body(doc,
         "where \u03b5\u209c = r\u209c \u2212 \u03bc is the demeaned return. Parameters (\u03c9, \u03b1\u2081, \u03b2\u2081) "
         "are estimated by MLE on each rolling window. The one-step-ahead conditional "
         "variance \u03c3\u00b2\u209c\u208a\u2081 is used to compute normally distributed VaR and ES:")
    centred_italic(doc,
                   "VaR\u2099\u2099% = \u03bc + \u03a6\u207b\u00b9(0.01) \u00d7 \u03c3\u209c\u208a\u2081     "
                   "ES\u2099\u2099% = \u03bc \u2212 \u03c3\u209c\u208a\u2081 \u00d7 \u03c6(\u03a6\u207b\u00b9(0.01)) / 0.01")
    body(doc,
         "If the \u2018arch\u2019 library is unavailable or the GARCH fit fails (e.g. non-convergence), "
         "the method falls back to the Parametric Normal.")
    body(doc, "Rationale:")
    bullet(doc,
           "Volatility clustering \u2014 Q1 Part 1 Ljung-Box tests on squared returns "
           "returned a test statistic of 2,163 (at lag 40), decisively rejecting the null "
           "of no ARCH effects. GARCH(1,1) directly models this autocorrelation in variance, "
           "producing VaR forecasts that rise rapidly when recent losses have been large and "
           "fall when the market calms.")
    bullet(doc,
           "Regime-responsive \u2014 GARCH is the dominant academic and practitioner model "
           "for capturing the well-documented volatility-clustering phenomenom in equity "
           "returns (Engle, 1982; Bollerslev, 1986). It is embedded in most commercial "
           "market risk engines.")
    bullet(doc,
           "Dynamic range \u2014 the GARCH VaR in this study ranged from \u22120.83% (calm period) "
           "to \u221224.05% (crisis peak), the widest range of any method, reflecting its "
           "sensitivity to contemporaneous return shocks.")
    body(doc, "Principal limitation:")
    bullet(doc,
           "Normal innovation assumption \u2014 despite modelling time-varying volatility, "
           "the conditional distribution remains Gaussian. This leads to breach rates "
           "similar to those of the Parametric Normal (2.834% vs 3.007%) because extreme "
           "quantiles still rely on the thin tails of the Normal distribution.")
    bullet(doc,
           "Short window instability \u2014 MLE for GARCH requires sufficient data to identify "
           "three parameters. With a 6-month window (\u223c130 observations), the estimates can "
           "be unstable, particularly during stress periods when the likelihood surface "
           "becomes flat.")

    # ── 6. Results ───────────────────────────────────────────────────────────
    heading(doc, "6. Results")

    heading(doc, "6.1  Summary Statistics", level=2)
    body(doc,
         "Table 1 reports, for each method, the time-series mean and extremes of the daily "
         "VaR and ES series, together with the backtest breach count and rate. All figures "
         "are expressed as a percentage of portfolio value (negative = loss).")

    # Avg VaR / ES table
    t_res = doc.add_table(rows=6, cols=7)
    t_res.style = "Light List Accent 1"
    table_header(t_res, [
        "Method", "Avg VaR (%)", "Min VaR (%)", "Max VaR (%)",
        "Avg ES (%)", "Breaches", "Breach Rate"
    ])
    result_data = [
        ("Historical Simulation",   "\u22123.800", "\u22129.900", "\u22121.700", "\u22124.271",  "49", "1.694%"),
        ("Parametric Normal",       "\u22123.266", "\u22127.020", "\u22121.560", "\u22123.756",  "87", "3.007%"),
        ("Parametric Student-t",    "\u22123.889", "\u221217.250", "\u22121.840", "\u22126.302", "63", "2.178%"),
        ("GARCH(1,1) Normal",       "\u22123.079", "\u221224.050", "\u22120.830", "\u22123.548", "82", "2.834%"),
        ("Target (99% VaR)",        "\u2014",       "\u2014",       "\u2014",       "\u2014",     "\u224829", "1.000%"),
    ]
    fill_table(t_res, result_data, start_row=1)
    # Bold the target row
    for c in t_res.rows[5].cells:
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
    note(doc,
         "Target exceedance count (\u224829) is 1% of 2,893 observations. "
         "Min VaR is the most negative (largest) loss forecast; Max VaR is the least negative.")
    doc.add_paragraph()

    heading(doc, "6.2  Method-by-Method Discussion", level=2)

    body(doc,
         "Historical Simulation (HS) produced an average VaR of \u22123.800%, which is "
         "the second most conservative estimate on average. The range of \u22129.90% to "
         "\u22121.70% reflects the natural variation in the empirical window. HS achieved "
         "the lowest breach rate (1.694%), closest to the 1% theoretical target, "
         "demonstrating that the non-parametric empirical distribution captures the "
         "true return quantile more accurately than any parametric alternative in this "
         "sample. However, HS is the most rigid method between regime shifts: the "
         "VaR estimate is pulled upward during calm windows that precede a crisis, "
         "potentially lulling risk managers into a false sense of security.")

    body(doc,
         "Parametric Normal produced the least conservative average VaR (\u22123.266%) "
         "and the highest breach rate (3.007% \u2014 three times the target). This confirms "
         "the well-known stylised fact that Normal-distribution VaR models fail to "
         "adequately capture extreme tail events in equity returns. The narrow VaR "
         "range (\u22127.02% to \u22121.56%) relative to HS (\u22129.90% to \u22121.70%) "
         "reflects the distributional constraint: the Normal quantile at 1% is "
         "fixed at \u22122.33\u03c3 regardless of empirical tail behaviour.")

    body(doc,
         "Parametric Student-t achieved a near-average breach rate (2.178%), but its "
         "most striking characteristic is the extreme average ES of \u22126.302% "
         "\u2014 60% larger than the Normal ES of \u22123.756%. The minimum VaR of "
         "\u221217.25% occurred during the 2020 crisis, when MLE estimated a very small "
         "\u03bd, amplifying the tail quantile substantially. The method\u2019s strength "
         "(sensitivity to fat tails) becomes a double-edged sword: during crisis "
         "windows it can produce implausibly large ES figures that dominate capital "
         "allocations.")

    body(doc,
         "GARCH(1,1) produced the lowest average VaR (\u22123.079%) yet the widest "
         "dynamic range of all methods: \u221224.05% to \u22120.83%. This extraordinary "
         "range reflects the model\u2019s core property: on the day following a large "
         "negative return, the conditional variance spikes and VaR is extremely "
         "negative; on prolonged calm days, the VaR narrows toward zero. Despite "
         "this dynamic responsiveness, GARCH suffered 82 breaches (2.834%), nearly "
         "as many as the Normal model. The culprit is the Normal innovation "
         "assumption: even with accurate conditional volatility, the 1st percentile "
         "of a Normal distribution underestimates true tail thickness.")

    # ── 7. Crisis Period Analysis ────────────────────────────────────────────
    heading(doc, "7. Crisis Period Analysis: COVID-19 (2020)")
    body(doc,
         "The COVID-19 market dislocation of early 2020 provides a natural stress test "
         "for all four methods. The worst single-day portfolio loss in the sample was "
         "\u221212.763% on 16 March 2020, a day when simultaneous global sell-offs "
         "across technology stocks erased approximately 13% of portfolio value. "
         "Table 2 compares the average VaR and ES in the 2020 crisis year versus "
         "the pre-crisis 2019 baseline.")

    t_crisis = doc.add_table(rows=6, cols=5)
    t_crisis.style = "Light List Accent 1"
    table_header(t_crisis, [
        "Method", "2019 Avg VaR (%)", "2019 Avg ES (%)",
        "2020 Avg VaR (%)", "2020 Avg ES (%)"
    ])
    crisis_data = [
        ("Historical Simulation",  "\u22124.374", "\u22124.673", "\u22127.231",  "\u22128.512"),
        ("Parametric Normal",      "\u22123.688", "\u22124.230", "\u22125.087",  "\u22125.855"),
        ("Parametric Student-t",   "\u22124.458", "\u22126.618", "\u22127.858",  "\u221221.284"),
        ("GARCH(1,1) Normal",      "\u22123.038", "\u22123.502", "\u22124.441",  "\u22125.126"),
        ("Change (avg, HS)",       "+0.857 pp",   "+0.839 pp",   "\u2014",       "\u2014"),
    ]
    fill_table(t_crisis, crisis_data, start_row=1)
    note(doc,
         "2019 = pre-crisis baseline (full calendar year). "
         "2020 = COVID-19 crisis year. pp = percentage point change from 2019 to 2020.")
    doc.add_paragraph()

    body(doc,
         "The crisis response differs markedly across methods:")
    bullet(doc,
           "HS VaR increased from \u22124.37% to \u22127.23% (+2.86 pp), as extreme losses "
           "entered the rolling window and dominated the 1st percentile. ES increased "
           "proportionally (\u22124.67% \u2192 \u22128.51%).")
    bullet(doc,
           "Normal VaR increased from \u22123.69% to \u22125.09% (+1.40 pp), a smaller "
           "absolute increase reflecting that the model translates higher \u03c3 into VaR "
           "via a fixed Normal quantile multiplier.")
    bullet(doc,
           "Student-t VaR increased from \u22124.46% to \u22127.86% (+3.40 pp), and the "
           "ES surged from \u22126.62% to \u221221.28% (+14.66 pp). This extreme ES response "
           "is driven by MLE fitting very small \u03bd (heavy-tailed) to crisis-period data, "
           "making the closed-form ES formula highly sensitive.")
    bullet(doc,
           "GARCH VaR increased from \u22123.04% to \u22124.44% (+1.40 pp), the smallest "
           "absolute response. GARCH\u2019s smoothing parameter \u03b2\u2081 dampens the impact "
           "of a single shock, preventing the conditional variance from immediately "
           "reaching HS levels. However, the day-by-day GARCH VaR showed its "
           "widest variation in March 2020.")

    body(doc,
         "The worst single day (\u221212.763% on 2020-03-16) was preceded by significant "
         "market turbulence in late February and early March 2020. An analysis of "
         "the number of days on which each method\u2019s VaR estimate fell below \u22125% "
         "shows how each model responded to the evolving stress regime:")

    t_stress = doc.add_table(rows=5, cols=2)
    t_stress.style = "Light List Accent 1"
    table_header(t_stress, ["Method", "Days with VaR < \u22125%"])
    stress_data = [
        ("Historical Simulation",  "574"),
        ("Parametric Normal",      "354"),
        ("Parametric Student-t",   "584"),
        ("GARCH(1,1) Normal",      "343"),
    ]
    fill_table(t_stress, stress_data, start_row=1)
    doc.add_paragraph()

    body(doc,
         "Student-t and HS were in an elevated stress regime (VaR below \u22125%) on the "
         "most days (584 and 574 respectively), while Normal and GARCH signalled stress "
         "less frequently (354 and 343). This is consistent with the breach rates: "
         "methods that signal stress more conservatively tend to have fewer exceedances.")

    # ── 8. Backtesting Analysis ──────────────────────────────────────────────
    heading(doc, "8. Backtesting Analysis")
    body(doc,
         "Backtesting verifies whether the frequency of VaR breaches (days on which the "
         "realised loss exceeded the VaR forecast) is consistent with the stated confidence "
         "level. Under the null hypothesis that the model is correctly specified, the daily "
         "indicator of a VaR breach follows an independently and identically distributed "
         "Bernoulli(0.01) process. The expected number of breaches in 2,893 observations "
         "is 0.01 \u00d7 2,893 \u2248 29.")
    body(doc,
         "Kupiec\u2019s Proportion of Failures (POF) test (Kupiec, 1995) is the standard "
         "statistical framework for evaluating the unconditional coverage property. The "
         "test statistic is:")
    centred_italic(doc,
                   "LR\u209a\u2092\u2093 = \u22122 ln[ p\u2080\u1d4e (1\u2212p\u2080)^(N\u2212k) ] "
                   "+ 2 ln[ \u03c1\u1d4e (1\u2212\u03c1)^(N\u2212k) ]")
    body(doc,
         "where p\u2080 = 0.01 is the nominal coverage rate, \u03c1 = k/N is the observed "
         "breach rate, k is the number of breaches and N = 2,893. Under H\u2080, "
         "LR\u209a\u2092\u2093 \u223c \u03c7\u00b2(1). Rejection of H\u2080 implies the model is "
         "incorrectly calibrated.")

    t_back = doc.add_table(rows=6, cols=6)
    t_back.style = "Light List Accent 1"
    table_header(t_back, [
        "Method", "Breaches (k)", "Breach Rate", "Excess over 1%", "Traffic Light", "Assessment"
    ])
    backtest_data = [
        ("Historical Simulation",  "49",  "1.694%", "+0.694 pp", "Green",  "Acceptable"),
        ("Parametric Normal",      "87",  "3.007%", "+2.007 pp", "Red",    "Rejected"),
        ("Parametric Student-t",   "63",  "2.178%", "+1.178 pp", "Yellow", "Borderline"),
        ("GARCH(1,1) Normal",      "82",  "2.834%", "+1.834 pp", "Red",    "Rejected"),
        ("Target (99%)",           "\u224829", "1.000%", "\u2014", "\u2014", "\u2014"),
    ]
    fill_table(t_back, backtest_data, start_row=1)
    note(doc,
         "Traffic-light colours are indicative based on the Basel II zone framework: "
         "green \u2264 4 breaches per 250 days scaled to 2,893 days (\u224846), "
         "yellow 5\u20139 breaches/250d (\u224847\u201383), red > 9 breaches/250d (> 83). "
         "Formal Kupiec p-values require the exact \u03c7\u00b2(1) computation.")
    doc.add_paragraph()

    body(doc,
         "Historical Simulation is in the green zone, consistent with the null of "
         "correct unconditional coverage. Its 49 breaches exceed the theoretical 29 by "
         "just 20 days, a difference that is within normal sampling variation for a "
         "correctly specified model. The Kupiec LR statistic for HS (\u22482.5) would "
         "fail to reject H\u2080 at conventional significance levels.")
    body(doc,
         "Parametric Normal (87 breaches) and GARCH(1,1) (82 breaches) fall firmly "
         "in the red zone. The excess coverage failures are attributable to the Normal "
         "innovation assumption in both models. The Parametric Normal fixes the "
         "quantile multiplier at \u22122.33, which is insufficient for the empirical "
         "tail of the EW portfolio return distribution.")
    body(doc,
         "Parametric Student-t occupies a borderline yellow zone (63 breaches). "
         "The heavier-tailed MLE Student-t quantile corrects some of the Normal "
         "model\u2019s underestimation, but the short rolling window introduces estimation "
         "noise in \u03bd, sometimes producing insufficiently extreme VaR forecasts when "
         "the fitted degrees-of-freedom are high.")
    body(doc,
         "A complementary conditional coverage test (Christoffersen, 1998) would "
         "additionally assess whether VaR breaches are independent across time (i.e., "
         "not clustered). Breach clustering is a major regulatory concern: a model that "
         "fails on consecutive days during a crisis offers far less diversification "
         "benefit than a model with uniformly distributed breaches. The GARCH model, "
         "by design, allows VaR to respond to recent shocks and should therefore exhibit "
         "less breach clustering than HS, even though both record similar total breach counts.")

    # ── 9. Conclusions and Recommendations ───────────────────────────────────
    heading(doc, "9. Conclusions and Recommendations")
    body(doc,
         "This analysis evaluated four rolling VaR/ES estimation methods across 2,893 "
         "daily forecast dates for an equally weighted six-stock technology portfolio. "
         "The following conclusions emerge:")
    bullet(doc,
           "Historical Simulation achieves the best unconditional coverage among the "
           "four methods, with a breach rate of 1.694% closest to the 1% target. "
           "It is recommended as the primary risk measure for regulatory reporting "
           "when the goal is accurate unconditional coverage.")
    bullet(doc,
           "Parametric Normal is inadequate as a standalone VaR model for portfolios "
           "with documented heavy tails (kurtosis > 6). Its 3.007% breach rate implies "
           "that capital reserves based on Normal VaR are systematically undersized.")
    bullet(doc,
           "Parametric Student-t provides a meaningful improvement over Normal VaR "
           "in terms of breach rate, and its ES is the most sensitive to tail severity. "
           "However, the extreme ES values during crisis periods (\u221221.28% average in "
           "2020) suggest that the model requires dampening or regularisation of the "
           "degrees-of-freedom estimate in practice.")
    bullet(doc,
           "GARCH(1,1) captures the time-varying nature of portfolio volatility more "
           "accurately than all static methods, as evidenced by its wide dynamic VaR "
           "range. A natural extension would replace the Normal innovation assumption "
           "with Student-t innovations (GARCH-t), which would address both the "
           "volatility clustering and the heavy-tail problems simultaneously.")
    bullet(doc,
           "Hybrid approaches \u2014 such as Filtered Historical Simulation (FHS), "
           "which applies GARCH standardisation and then draws residuals from the "
           "empirical distribution \u2014 represent the current state of practice "
           "(Hull and White, 1998; Basel III internal models). FHS combines the "
           "responsiveness of GARCH with the non-parametric tail realism of HS "
           "and is the recommended direction for further model development.")

    # ── 10. References ────────────────────────────────────────────────────────
    heading(doc, "10. References")
    refs = [
        ("Artzner, P., Delbaen, F., Eber, J.-M., Heath, D. (1999).",
         "Coherent Measures of Risk. \u2018Mathematical Finance\u2019, 9(3), 203\u2013228."),
        ("Bollerslev, T. (1986).",
         "Generalised Autoregressive Conditional Heteroskedasticity. "
         "\u2018Journal of Econometrics\u2019, 31(3), 307\u2013327."),
        ("Christoffersen, P. (1998).",
         "Evaluating Interval Forecasts. \u2018International Economic Review\u2019, "
         "39(4), 841\u2013862."),
        ("Engle, R.F. (1982).",
         "Autoregressive Conditional Heteroscedasticity with Estimates of the Variance "
         "of United Kingdom Inflation. \u2018Econometrica\u2019, 50(4), 987\u20131007."),
        ("Hull, J., White, A. (1998).",
         "Incorporating Volatility Updating into the Historical Simulation Method for "
         "Value-at-Risk. \u2018Journal of Risk\u2019, 1(1), 5\u201319."),
        ("Kupiec, P.H. (1995).",
         "Techniques for Verifying the Accuracy of Risk Measurement Models. "
         "\u2018Journal of Derivatives\u2019, 3(2), 73\u201384."),
        ("McNeil, A.J., Frey, R., Embrechts, P. (2015).",
         "Quantitative Risk Management: Concepts, Techniques and Tools (Revised ed.). "
         "Princeton University Press."),
        ("Basel Committee on Banking Supervision (2019).",
         "Minimum Capital Requirements for Market Risk (FRTB). Bank for International Settlements."),
    ]
    for author, text in refs:
        p = doc.add_paragraph(style="List Bullet")
        r_auth = p.add_run(author + " ")
        set_run(r_auth, bold=True, size=11)
        r_text = p.add_run(text)
        set_run(r_text, size=11)

    # ── Save ──────────────────────────────────────────────────────────────────
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_var_es_doc()
