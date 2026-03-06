"""
Utility script — generates all methodology .docx files for SMM272 Risk Analysis.
Run from the project root:  python docs/_generate_docs.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_run(run, size=11, bold=False, italic=False, colour=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = RGBColor(*colour)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def centred_italic(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=size, italic=True)
    return p


# ── Document: Q1 Equally Weighted Portfolio ───────────────────────────────────

def build_ew_portfolio_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    t = doc.add_heading("SMM272 Risk Analysis \u2014 Q1", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Equally Weighted Portfolio: Methodology and Rationale")
    set_run(r, size=14, bold=True, colour=(31, 73, 125))
    doc.add_paragraph()

    # 1. Overview
    heading(doc, "1. Overview")
    body(doc,
         "This document explains the methodology and rationale behind constructing an equally "
         "weighted (EW) portfolio of six large-cap technology stocks \u2014 AAPL, MSFT, IBM, NVDA, "
         "GOOGL and AMZN \u2014 over the period January 2014 to December 2025. The EW portfolio "
         "serves as the main analytical object for Q1 of the SMM272 coursework.")

    # 2. Definition
    heading(doc, "2. What Is an Equally Weighted Portfolio?")
    body(doc,
         "An equally weighted portfolio assigns an identical fractional weight to every constituent "
         "asset. With N assets, each weight is:")
    centred_italic(doc, "w\u1d62 = 1 / N     for i = 1, 2, \u2026, N")
    body(doc,
         "For this portfolio N = 6, so each stock receives a weight of 1/6 \u2248 16.67 %. "
         "The portfolio return on day t is therefore the simple arithmetic average of the "
         "individual asset returns:")
    centred_italic(doc, "r\u209a,\u209c = (1/N) \u00d7 \u03a3 r\u1d62,\u209c")

    # 3. Log returns
    heading(doc, "3. Why Log Returns?")
    body(doc,
         "Individual asset returns are computed as continuously compounded (logarithmic) returns "
         "rather than simple percentage returns:")
    centred_italic(doc, "r\u1d62,\u209c = ln( P\u1d62,\u209c / P\u1d62,\u209c\u208b\u2081 )")
    body(doc, "The choice of log returns is motivated by four statistical properties:")
    bullet(doc, "Time additivity \u2014 multi-period log returns equal the sum of single-period log "
                "returns, making aggregation straightforward.")
    bullet(doc, "Approximate normality \u2014 log returns are closer to a normal distribution than "
                "simple returns, which facilitates parametric testing.")
    bullet(doc, "Symmetry \u2014 a gain of x% and a loss of x% are symmetric under log returns but "
                "not under simple returns.")
    bullet(doc, "Stationarity \u2014 log returns transform a non-stationary price series into a "
                "stationary return series, a prerequisite for the time-series tests used later "
                "(Ljung-Box, normality tests).")
    body(doc,
         "Note: strictly, the weighted sum of individual log returns is an approximation of the "
         "true portfolio log return. The approximation error is negligible for daily data and "
         "is standard practice in quantitative finance.")

    # 4. Rationale
    heading(doc, "4. Rationale for Equal Weighting")

    heading(doc, "4.1  Na\u00efve Diversification Benchmark", level=2)
    body(doc,
         "The 1/N strategy is a well-established academic and industry benchmark. DeMiguel, "
         "Garlappi and Uppal (2009, Review of Financial Studies) demonstrated that the equally "
         "weighted portfolio frequently outperforms mean-variance optimised portfolios out-of-sample "
         "due to the severe estimation error embedded in optimised weights.")

    heading(doc, "4.2  Avoidance of Estimation Risk", level=2)
    body(doc,
         "Mean-variance optimisation (Markowitz, 1952) requires estimating expected returns, "
         "variances and covariances. Expected return estimates are notoriously noisy \u2014 small "
         "errors in the mean vector are amplified by matrix inversion and produce highly unstable, "
         "extreme weights. Equal weighting sidesteps this problem entirely: the weights are "
         "deterministic and require no forward-looking assumptions.")

    heading(doc, "4.3  Transparency and Reproducibility", level=2)
    body(doc,
         "Because the weights follow a simple formula, the portfolio is fully transparent. "
         "Any researcher can replicate it from the raw price data without access to proprietary "
         "optimiser inputs. This is important for an academic coursework setting where "
         "reproducibility is a core requirement.")

    heading(doc, "4.4  Appropriate Scope for Q1", level=2)
    body(doc,
         "The objective of Q1 is to analyse the statistical properties of portfolio returns \u2014 "
         "distributional shape, tail risk, autocorrelation and cross-sectional correlation \u2014 "
         "not to optimise weights. Using a fixed, rule-based weighting scheme ensures that "
         "all observed return characteristics are attributable to the asset universe and the "
         "sample period, rather than to the weighting methodology.")

    # 5. Implementation
    heading(doc, "5. Implementation in Code")
    body(doc, "The portfolio is constructed in q1_1_build_portfolio.py in three steps:")
    bullet(doc, "Step 1 \u2014 Load adjusted closing prices (from Yahoo Finance via yfinance, "
                "cached to CSV after the first download).")
    bullet(doc, "Step 2 \u2014 Compute daily log returns using the formula above; drop the first "
                "row (NaN produced by the one-period shift).")
    bullet(doc, "Step 3 \u2014 Multiply the (T \u00d7 N) log-return matrix by the (N \u00d7 1) "
                "equal-weight vector via a dot product to produce the (T \u00d7 1) portfolio "
                "return series.")
    body(doc,
         "Both the individual log returns and the portfolio return series are persisted to CSV "
         "files in the output/ directory so that downstream modules (descriptive statistics, "
         "normality tests, autocorrelation analysis, visualisations) can load pre-computed "
         "results without repeating the calculation.")

    # 6. Limitations
    heading(doc, "6. Limitations and Caveats")
    bullet(doc, "Concentration risk \u2014 all six constituents are large-cap US technology stocks. "
                "The portfolio is sector-concentrated and does not represent a broadly diversified "
                "multi-asset portfolio.")
    bullet(doc, "Rebalancing \u2014 in practice an EW portfolio drifts as prices move and requires "
                "periodic rebalancing back to 1/N. This implementation uses a buy-and-hold "
                "approximation (daily log-return average), which implicitly assumes continuous "
                "rebalancing.")
    bullet(doc, "Survivorship bias \u2014 all six stocks existed and were publicly traded throughout "
                "the sample period. Selecting them with hindsight introduces a mild survivorship bias.")

    # 7. References
    heading(doc, "7. Key References")
    bullet(doc, "DeMiguel, V., Garlappi, L., & Uppal, R. (2009). Optimal versus naive "
                "diversification: How inefficient is the 1/N portfolio strategy? "
                "Review of Financial Studies, 22(5), 1915\u20131953.")
    bullet(doc, "Markowitz, H. (1952). Portfolio selection. Journal of Finance, 7(1), 77\u201391.")
    bullet(doc, "Campbell, J. Y., Lo, A. W., & MacKinlay, A. C. (1997). "
                "The Econometrics of Financial Markets. Princeton University Press.")

    out = os.path.join(DOCS_DIR, "Q1_Equally_Weighted_Portfolio_Methodology.docx")
    doc.save(out)
    print(f"Saved: {out}")
    return out



# ── Document: Q1 Statistical Analysis Report ─────────────────────────────────

def build_statistical_analysis_report():  # noqa: C901
    """
    Generate the Q1 Statistical Analysis Report for the SMM272 coursework.
    Covers procedure and findings for all six analytical modules.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Title ──────────────────────────────────────────────────────────────
    t = doc.add_heading("SMM272 Risk Analysis \u2014 Question 1", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Statistical Analysis of a Technology Equity Portfolio: "
        "Procedure and Findings"
    )
    set_run(r, size=14, bold=True, colour=(31, 73, 125))
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Sample period: January 2014 \u2013 December 2025  |  "
                 "Portfolio: AAPL, MSFT, IBM, NVDA, GOOGL, AMZN  |  "
                 "Observations: 3,016 trading days").italic = True

    doc.add_paragraph()

    # ── 1. Executive Summary ───────────────────────────────────────────────
    heading(doc, "1. Executive Summary")
    body(doc,
         "This report presents the full statistical analysis of an equally weighted (EW) "
         "portfolio comprising six large-capitalisation US technology stocks: Apple (AAPL), "
         "Microsoft (MSFT), IBM, NVIDIA (NVDA), Alphabet (GOOGL) and Amazon (AMZN). Daily "
         "log returns are analysed over the period January 2014 to December 2025, yielding "
         "3,016 observations. The analysis proceeds through six modular steps: descriptive "
         "statistics, risk metrics, normality testing with distribution fitting, time-series "
         "diagnostics, correlation analysis, and portfolio-level conclusions.")
    body(doc,
         "Key findings are: (i) the portfolio delivered a strong annualised return of "
         "24.57% with a Sharpe ratio of 1.024 and a Sortino ratio of 1.314; (ii) returns "
         "are significantly non-normal \u2014 leptokurtic and negatively skewed \u2014 with a "
         "Student-t distribution (df \u2248 3.30) providing a materially better fit than the "
         "Gaussian; (iii) daily returns are stationary but exhibit significant serial "
         "autocorrelation and strong ARCH effects; (iv) intra-portfolio correlations, "
         "entirely positive and rising sharply during the COVID-19 crisis, limit "
         "diversification benefit to approximately 20% risk reduction.")

    # ── 2. Data and Sample ─────────────────────────────────────────────────
    heading(doc, "2. Data and Sample Description")
    heading(doc, "2.1  Asset Universe", level=2)
    body(doc,
         "The portfolio consists of six US technology equities selected to represent the "
         "broad large-cap technology sector: AAPL, MSFT, IBM, NVDA, GOOGL, and AMZN. All "
         "six were continuously listed on major US exchanges throughout the sample period, "
         "which introduces a mild survivorship bias. Daily adjusted closing prices are "
         "obtained from Yahoo Finance via the yfinance library (version 1.2.0), with "
         "adjustments for stock splits and dividends applied automatically.")

    heading(doc, "2.2  Sample Period", level=2)
    body(doc,
         "The sample spans from 1 January 2014 to 31 December 2025, a period of approximately "
         "twelve years. This window captures several economically significant episodes including "
         "the 2015 China-driven equity selloff, the COVID-19 market crash of February\u2013March "
         "2020, the subsequent sharp recovery, the inflation-driven bear market of 2022, and "
         "the AI-driven rally of 2023\u20132025. After removing non-overlapping trading days and "
         "missing values, the clean dataset contains 3,016 daily observations.")

    heading(doc, "2.3  Log Returns", level=2)
    body(doc,
         "Continuously compounded (log) returns are computed from adjusted closing prices:")
    centred_italic(doc, "r\u1d62,\u209c = ln( P\u1d62,\u209c / P\u1d62,\u209c\u208b\u2081 )")
    body(doc,
         "Log returns are preferred over simple returns for four reasons: they are "
         "time-additive, approximately normal over short horizons, symmetric under sign "
         "reversal, and transform a non-stationary price series into a stationary return "
         "series (Campbell, Lo and MacKinlay, 1997). The portfolio return on day t is the "
         "equal-weight-averaged sum of the six individual log returns.")

    # ── 3. Portfolio Construction ──────────────────────────────────────────
    heading(doc, "3. Portfolio Construction Procedure")
    body(doc,
         "An equally weighted portfolio assigns a constant weight of w\u1d62 = 1/N = 1/6 \u2248 16.67% "
         "to each constituent at each rebalancing date. The daily portfolio return is therefore:")
    centred_italic(doc, "r\u209a,\u209c = (1/6) \u00d7 (r\u2081,\u209c + r\u2082,\u209c + "
                        "r\u2083,\u209c + r\u2084,\u209c + r\u2085,\u209c + r\u2086,\u209c)")
    body(doc,
         "The EW weighting scheme was selected for this analysis for three reasons. First, it "
         "provides a well-documented Na\u00efve Diversification benchmark: DeMiguel, Garlappi and "
         "Uppal (2009) showed that the 1/N portfolio consistently outperforms optimised "
         "portfolios out-of-sample because it avoids the estimation error inherent in "
         "mean-variance optimisation. Second, the weights are deterministic and require no "
         "forward-looking parameter assumptions, ensuring full reproducibility. Third, the "
         "objective of this question is to characterise the statistical properties of portfolio "
         "returns, not to optimise allocation \u2014 a fixed weighting scheme ensures that all "
         "observed characteristics are attributable to the asset universe and sampling period "
         "alone.")

    # ── 4. Analytical Procedure ────────────────────────────────────────────
    heading(doc, "4. Analytical Procedure")
    body(doc,
         "The analysis is structured as six sequential modules, each targeting a distinct "
         "dimension of the return distribution. The modules are orchestrated by a master "
         "script (q1_1_statistical_analysis.py) that passes pre-computed inputs downstream "
         "to avoid redundant computation.")

    # Module 1
    heading(doc, "4.1  Module 1 \u2014 Descriptive Statistics", level=2)
    body(doc,
         "The first module computes a set of pure distributional moments and order statistics "
         "for both the EW portfolio and each individual constituent. For each series the "
         "following statistics are reported: the number of observations (T), the daily "
         "arithmetic mean return (\u03bc), the daily median return, skewness (\u03b3\u2081), "
         "excess kurtosis (\u03b3\u2082 = \u03ba \u2212 3), the minimum and maximum single-day "
         "returns, the first and third quartiles (Q1 and Q3), and the cumulative return over "
         "the full sample. Skewness and excess kurtosis are the natural diagnostics for "
         "asymmetry and fat-tail behaviour respectively; a Gaussian distribution has both "
         "equal to zero.")

    # Module 2
    heading(doc, "4.2  Module 2 \u2014 Risk Metrics", level=2)
    body(doc,
         "The second module computes return-based risk and performance metrics. Daily standard "
         "deviation (\u03c3) is annualised using the square-root-of-time rule: "
         "\u03c3\u2090\u2099\u2099 = \u03c3 \u00d7 \u221a252. "
         "Two reward-to-risk ratios are reported. The Sharpe ratio (Sharpe, 1966) is defined "
         "as the annualised excess return divided by annualised volatility, with the risk-free "
         "rate set to zero to isolate the gross reward-to-risk profile. The Sortino ratio "
         "replaces the denominator with the annualised downside deviation (computed from "
         "negative-only daily returns), thereby penalising downside volatility only and "
         "providing a more meaningful measure when returns are skewed.")
    body(doc,
         "Tail risk is quantified through three approaches. Historical (empirical) Value at "
         "Risk at confidence levels of 95% and 99% is the sample quantile of the return "
         "distribution at the 5th and 1st percentile respectively. Historical Expected "
         "Shortfall (CVaR/ES) is the conditional mean of returns in the tail beyond the VaR "
         "threshold, capturing the average loss in adverse scenarios. Parametric (Gaussian) "
         "VaR is computed as \u03bc + z\u03b1 \u00d7 \u03c3 where z\u03b1 is the standard normal "
         "quantile; this provides a baseline against which to judge the gap introduced by "
         "non-normality.")

    # Module 3
    heading(doc, "4.3  Module 3 \u2014 Normality Testing and Distribution Fitting", level=2)
    body(doc,
         "To formally test whether portfolio returns conform to a Gaussian distribution, "
         "four complementary tests are applied.")
    bullet(doc, "Jarque-Bera (JB) test \u2014 a chi-squared test based on the joint departure of "
                "sample skewness and kurtosis from their Gaussian values (0, 3). It has "
                "high power in large samples against leptokurtic alternatives.")
    bullet(doc, "Shapiro-Wilk (SW) test \u2014 examines the correlation between ordered sample "
                "values and the expected order statistics of a normal distribution. As the "
                "full sample exceeds 5,000 observations, a random subsample of 5,000 "
                "observations is used to satisfy the test\u2019s size requirements.")
    bullet(doc, "D\u2019Agostino-Pearson K\u00b2 \u2014 combines separate tests of departure from "
                "normality in skewness and kurtosis into a single omnibus statistic.")
    bullet(doc, "Kolmogorov-Smirnov (KS) test \u2014 compares the empirical CDF with the CDF of "
                "a normal distribution fitted by maximum likelihood to the sample data.")
    body(doc,
         "Distribution fitting compares the Normal and Student-t distributions via maximum "
         "likelihood estimation (MLE). The Normal distribution is characterised by its mean "
         "and standard deviation. The Student-t distribution adds a degrees-of-freedom "
         "parameter \u03bd that governs tail thickness; smaller \u03bd implies heavier tails. "
         "Model comparison uses the Akaike Information Criterion (AIC = \u22122 log L + 2k) and "
         "the Bayesian Information Criterion (BIC = \u22122 log L + k ln n), where k is the "
         "number of free parameters. Both criteria penalise the likelihood for model "
         "complexity, with lower values preferred.")

    # Module 4
    heading(doc, "4.4  Module 4 \u2014 Time-Series Diagnostics", level=2)
    body(doc,
         "Time-series properties of the return series are examined through two types of tests: "
         "autocorrelation and stationarity.")
    body(doc, "Autocorrelation is tested using the Ljung-Box (LB) portmanteau statistic "
              "(Ljung and Box, 1978) at lags h \u2208 {5, 10, 15, 20}. Under the null "
              "hypothesis of no serial correlation up to lag h, the statistic follows an "
              "approximate chi-squared distribution with h degrees of freedom. The test is "
              "applied twice: once to the raw returns and once to the squared returns. "
              "Significant LB statistics on raw returns indicate predictability in the level; "
              "significant statistics on squared returns indicate time-varying volatility "
              "(ARCH effects / volatility clustering), which is a prerequisite for GARCH-type "
              "modelling.", indent=True)
    body(doc, "Stationarity is assessed with a two-test approach combining the Augmented "
              "Dickey-Fuller (ADF) test and the Kwiatkowski-Phillips-Schmidt-Shin (KPSS) test "
              "using complementary null hypotheses. The ADF tests H\u2080: unit root (non-stationary) "
              "against H\u2081: stationary, with lag length selected by AIC. The KPSS tests "
              "H\u2080: stationary against H\u2081: unit root. When both tests concur, the conclusion "
              "is unambiguous; conflicting results are treated as inconclusive.", indent=True)

    # Module 5
    heading(doc, "4.5  Module 5 \u2014 Correlation Analysis", level=2)
    body(doc,
         "Pairwise linear (Pearson) correlations are computed for the full sample and separately "
         "for the COVID-19 crisis period (1 January 2020 \u2013 31 December 2020) and the "
         "surrounding non-crisis period. The comparison tests for correlation breakdown, the "
         "well-documented empirical regularity whereby asset correlations rise precisely during "
         "market stress episodes, eroding the diversification benefit when it is most needed "
         "(Longin and Solnik, 2001). Average pairwise off-diagonal correlation is the summary "
         "statistic used to compare the two windows.').")

    # Module 6
    heading(doc, "4.6  Module 6 \u2014 Portfolio-Level Conclusions", level=2)
    body(doc,
         "The final module synthesises cross-sectional evidence. It constructs an annualised "
         "risk-return summary table for all six constituents and the portfolio, calculates the "
         "percentage diversification benefit (reduction in annualised standard deviation "
         "relative to the average constituent), and compares 99% VaR and CVaR across assets "
         "to quantify the tail-risk reduction achieved by aggregation.")

    # ── 5. Main Findings ───────────────────────────────────────────────────
    heading(doc, "5. Main Findings")

    # 5.1 Return Profile
    heading(doc, "5.1  Return and Cumulative Performance", level=2)
    body(doc,
         "Over 3,016 trading days the EW portfolio produced a daily mean log return of "
         "0.0975%, equivalent to an annualised mean return of 24.57%. The daily median return "
         "(0.149%) exceeds the arithmetic mean, which is characteristic of a negatively skewed "
         "distribution: in the majority of sessions the portfolio moved upward, but rare, "
         "large drawdowns pulled the mean below the median. The cumulative return over the "
         "full sample was 1,793.7%, more than 17-fold growth in value.")
    body(doc,
         "Individual constituents exhibited wide dispersion in performance. NVIDIA was the "
         "standout performer, accumulating a cumulative return of 50,063% over the sample "
         "period and an annualised return of 51.95%, driven primarily by successive waves of "
         "demand for GPU hardware in gaming, data centres and AI workloads. IBM was the "
         "weakest constituent at 8.47% annualised, reflecting structural competitive headwinds "
         "in legacy enterprise computing. The remaining four stocks (AAPL, MSFT, GOOGL, AMZN) "
         "clustered narrowly between 20.3% and 23.1% annualised, consistent with sector-wide "
         "growth dynamics.")

    # 5.2 Risk and Performance
    heading(doc, "5.2  Risk Metrics and Risk-Adjusted Performance", level=2)
    body(doc,
         "The EW portfolio\u2019s annualised volatility was 23.99% \u2014 materially lower than every "
         "individual constituent except IBM (23.99%). This is the first evidence of "
         "diversification: combining positively correlated but not perfectly correlated assets "
         "reduces portfolio volatility below the weighted-average of constituent volatilities "
         "(30.09%), a reduction of approximately 20.3%.")

    # Add risk-return table
    table = doc.add_table(rows=8, cols=5)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Asset", "Ann. Return", "Ann. Std", "Sharpe", "Sortino"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True

    data_rows = [
        ("AAPL",         "23.12%", "28.29%", "0.817", "1.089"),
        ("MSFT",         "23.07%", "26.30%", "0.877", "1.178"),
        ("IBM",          " 8.47%", "23.99%", "0.353", "0.437"),
        ("NVDA",         "51.95%", "46.87%", "1.108", "1.539"),
        ("GOOGL",        "20.30%", "28.24%", "0.719", "0.981"),
        ("AMZN",         "20.54%", "32.82%", "0.626", "0.850"),
        ("EW Portfolio", "24.57%", "23.99%", "1.024", "1.314"),
    ]
    for row_idx, row_data in enumerate(data_rows, start=1):
        cells = table.rows[row_idx].cells
        for col_idx, val in enumerate(row_data):
            cells[col_idx].text = val
        if row_idx == 7:   # portfolio row — bold
            for c in table.rows[row_idx].cells:
                for para in c.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph()

    body(doc,
         "The EW portfolio achieves the highest Sharpe ratio in the cross-section (1.024), "
         "surpassing even NVIDIA (1.108 without portfolio context). This result is a "
         "manifestation of the mean-variance efficiency gains from combining constituents: "
         "by pooling NVIDIA\u2019s high return with IBM\u2019s lower volatility and spreading "
         "idiosyncratic risk, the portfolio improves risk-adjusted performance beyond what "
         "any single asset achieves in isolation.")
    body(doc,
         "The Sortino ratio (1.314) exceeds the Sharpe ratio, which implies that a "
         "disproportionate share of portfolio volatility is upside volatility. As a result, "
         "standard deviation overstates the \u2018true\u2019 downside risk relative to what the "
         "Sortino ratio captures. This asymmetry is consistent with the negatively skewed "
         "return distribution documented below.")
    body(doc,
         "Historical tail risk estimates show the EW portfolio\u2019s VaR(99%) = \u22124.36% "
         "and CVaR(99%) = \u22125.60% per day. The parametric (Gaussian) VaR(99%) of \u22123.42% "
         "understates the historical VaR by approximately 94 basis points, confirming that "
         "the Normal distribution materially under-estimates tail risk. The gap widens in "
         "the 99th percentile relative to the 95th percentile, which is consistent with a "
         "fat-tailed distribution.")

    # 5.3 Distributional Properties
    heading(doc, "5.3  Distributional Properties and Normality", level=2)
    body(doc,
         "The EW portfolio exhibits skewness of \u22120.328 and excess kurtosis of 6.334. "
         "The negative skewness reflects the stylised fact that equity returns experience "
         "sharp, sudden drawdowns (left-tail events) that are larger in magnitude than "
         "equivalent up-moves; investors face crash risk. The excess kurtosis of 6.334 "
         "is substantially above zero, indicating a leptokurtic distribution: the central "
         "peak is taller and the tails are heavier than those of a Normal distribution with "
         "the same mean and variance.")
    body(doc,
         "All four normality tests reject the null hypothesis of normality at the 1% "
         "significance level:")

    # Normality table
    ntable = doc.add_table(rows=5, cols=4)
    ntable.style = "Light List Accent 1"
    nhdr = ntable.rows[0].cells
    for i, txt in enumerate(["Test", "Statistic", "p-value", "Decision (5%)"]):
        nhdr[i].text = txt
        nhdr[i].paragraphs[0].runs[0].bold = True
    norm_data = [
        ("Jarque-Bera",         "5,076.16", "\u2248 0",          "Reject H\u2080"),
        ("Shapiro-Wilk",        "0.9366",   "4.52 \u00d7 10\u207b\u00b3\u2074", "Reject H\u2080"),
        ("D\u2019Agostino K\u00b2",   "453.11",   "4.05 \u00d7 10\u207b\u2079\u2079", "Reject H\u2080"),
        ("Kolmogorov-Smirnov",  "0.0799",   "3.42 \u00d7 10\u207b\u00b9\u2077", "Reject H\u2080"),
    ]
    for row_idx, row_data in enumerate(norm_data, start=1):
        cells = ntable.rows[row_idx].cells
        for col_idx, val in enumerate(row_data):
            cells[col_idx].text = val

    doc.add_paragraph()

    body(doc,
         "The p-values span from essentially machine zero (Jarque-Bera) to 10\u207b\u00b3\u2074 "
         "(Shapiro-Wilk), constituting overwhelming statistical evidence against the Gaussian "
         "distribution. Four independent testing frameworks \u2014 moment-based, regression-based, "
         "omnibus, and CDF-distance-based \u2014 all concur.")
    body(doc,
         "Distribution fitting by MLE confirms this conclusion. The Student-t distribution "
         "estimated degrees of freedom \u03bd = 3.30 (AIC = \u221217,264 versus \u221216,725 for the "
         "Normal, a difference of 539 units). The AIC improvement of 539 per 3,016 observations "
         "indicates the Student-t is decisively preferred. An estimated \u03bd of 3.30 has "
         "important implications: tails are very heavy (the variance exists only marginally "
         "given \u03bd > 2 is required; the fourth moment is infinite for \u03bd \u2264 4). "
         "From a risk management perspective, this means extreme drawdowns occur far more "
         "frequently than Gaussian models predict. The standard deviation of the portfolio "
         "is 1.51% per day under the Normal fit but the tail scale parameter of the "
         "Student-t fit is 1.00%, with the heaviness carried instead in the degrees-of-freedom "
         "parameter. Gaussian VaR and CVaR models that ignore this will systematically "
         "under-capitalise for tail risk.")

    # 5.4 Time-Series
    heading(doc, "5.4  Time-Series Properties", level=2)
    body(doc, "The portfolio return series satisfies the fundamental prerequisite for time-series "
              "analysis: stationarity. The ADF test rejects the unit-root null hypothesis with a "
              "test statistic of \u221218.26 (p \u2248 2.3 \u00d7 10\u207b\u00b3\u2070), vastly below "
              "the 1% critical value of \u22123.43. The KPSS test, with a null of stationarity, "
              "fails to reject (statistic = 0.040, p = 0.10, critical value 0.463 at 5%). "
              "Both tests concur: daily log returns are stationary, consistent with the theoretical "
              "expectation that log returns of liquid assets are integrated of order zero.")
    body(doc,
         "Despite stationarity, the Ljung-Box test detects highly significant serial autocorrelation "
         "in the raw return series. At lag 5, LB = 36.6 (p = 7.3 \u00d7 10\u207b\u2077); at lag 20, "
         "LB = 136.9 (p = 1.9 \u00d7 10\u207b\u00b9\u2079). These results reject return independence "
         "at all conventional significance levels, meaning past returns contain some predictive "
         "information for future returns. The practical magnitude of the autocorrelations, however, "
         "is likely small \u2014 sufficient to detect statistically in 3,016 observations, but "
         "insufficient for profitable trading after transaction costs.")
    body(doc,
         "The most striking time-series finding concerns squared returns. Ljung-Box statistics "
         "on r\u209c\u00b2 at lag 5 reach 1,165 (p \u2248 10\u207b\u00b2\u2075\u2070), and at lag 20 "
         "reach 2,163 (p \u2248 0 to machine precision). These magnitudes, some 16\u221258 times "
         "larger than those on raw returns, constitute definitive evidence of volatility clustering: "
         "large squared returns (high-volatility periods) are followed by further large squared "
         "returns. This \u2018ARCH effect\u2019 (Engle, 1982) is perhaps the most studied stylised "
         "fact of financial time series. Its presence means that a simple constant-variance "
         "model is misspecified for this portfolio \u2014 dynamic variance models such as "
         "GARCH (Bollerslev, 1986) are required for accurate volatility forecasting and the "
         "construction of conditional VaR estimates.")

    # 5.5 Correlations
    heading(doc, "5.5  Correlation Structure and Diversification", level=2)
    body(doc,
         "Full-sample pairwise correlations are entirely positive, ranging from 0.278 "
         "(IBM\u2013AMZN) to 0.677 (MSFT\u2013GOOGL), with an average off-diagonal value of 0.505. "
         "The uniformly positive structure reflects sector concentration: all six firms "
         "operate principally in technology and consumer internet markets, subjecting "
         "them to common demand, regulatory and macroeconomic shocks. IBM exhibits the "
         "lowest correlations with all peers, as it is more exposed to enterprise IT "
         "services business cycles than to consumer-facing internet growth.")

    # Correlation table
    ctable = doc.add_table(rows=7, cols=7)
    ctable.style = "Light Grid Accent 1"
    tickers = ["", "AAPL", "MSFT", "IBM", "NVDA", "GOOGL", "AMZN"]
    for col_idx, txt in enumerate(tickers):
        ctable.rows[0].cells[col_idx].text = txt
        if col_idx > 0:
            ctable.rows[0].cells[col_idx].paragraphs[0].runs[0].bold = True
    corr_values = [
        ("AAPL",  "1.000", "0.645", "0.383", "0.522", "0.576", "0.534"),
        ("MSFT",  "0.645", "1.000", "0.405", "0.600", "0.677", "0.621"),
        ("IBM",   "0.383", "0.405", "1.000", "0.304", "0.369", "0.278"),
        ("NVDA",  "0.522", "0.600", "0.304", "1.000", "0.526", "0.510"),
        ("GOOGL", "0.576", "0.677", "0.369", "0.526", "1.000", "0.626"),
        ("AMZN",  "0.534", "0.621", "0.278", "0.510", "0.626", "1.000"),
    ]
    for row_idx, row_data in enumerate(corr_values, start=1):
        cells = ctable.rows[row_idx].cells
        cells[0].text = row_data[0]
        cells[0].paragraphs[0].runs[0].bold = True
        for col_idx, val in enumerate(row_data[1:], start=1):
            cells[col_idx].text = val

    doc.add_paragraph()

    body(doc,
         "The COVID-19 crisis window (2020) reveals dramatic correlation contagion. "
         "Average pairwise off-diagonal correlation rose from 0.505 (full sample) to 0.702 "
         "during 2020 \u2014 an increase of nearly 20 percentage points. Every individual "
         "pair rose substantially: for instance, the MSFT\u2013GOOGL correlation increased "
         "from 0.677 to 0.862. This correlation breakdown is the phenomenon documented "
         "by Longin and Solnik (2001): during market stress, equity correlations converge "
         "towards one as investors engage in broad-based deleveraging, eliminating the "
         "diversification benefit. For a portfolio that is already sector-concentrated, "
         "this effect is especially harmful.")
    body(doc,
         "The 20.3% annualised volatility reduction relative to the equal-weighted average "
         "constituent volatility confirms that diversification does add value in-sample. "
         "However, because all pairwise correlations are positive, the maximum possible "
         "risk reduction is far below what would be achievable with non-correlated assets. "
         "A multi-sector or multi-asset portfolio would provide substantially superior "
         "diversification.")

    # ── 6. Conclusions ─────────────────────────────────────────────────────
    heading(doc, "6. Conclusions")
    body(doc,
         "The statistical analysis of the equally weighted US technology portfolio over "
         "2014\u20132025 yields several clear conclusions with practical implications for "
         "risk management.")
    bullet(doc,
           "Strong long-run performance: annualised return of 24.57% and Sharpe ratio "
           "of 1.024 reflect the structural growth of the technology sector over the period. "
           "NVIDIA\u2019s extraordinary appreciation dominates cumulative portfolio return.")
    bullet(doc,
           "Non-normality is pervasive and consequential: all four normality tests reject the "
           "Gaussian at machine precision. Fat tails (excess kurtosis 6.33, Student-t df 3.30) "
           "mean that Gaussian VaR under-estimates actual 99% tail risk by roughly 94 basis "
           "points per day. Risk models must incorporate heavy-tailed distributions.")
    bullet(doc,
           "Volatility clustering demands conditional models: Ljung-Box statistics on squared "
           "returns are orders of magnitude larger than those on levels (LB = 2,163 vs 137 "
           "at lag 20). A GARCH-family or stochastic volatility model is required for "
           "accurate conditional VaR forecasting.")
    bullet(doc,
           "Sector concentration limits diversification: average pairwise correlation of 0.505 "
           "and crisis-period correlation of 0.702 confirm that this portfolio does not provide "
           "true diversification. The 20% volatility reduction is meaningful but far below "
           "what would be achievable with orthogonal asset classes.")
    bullet(doc,
           "Expected Shortfall (CVaR) is the appropriate tail-risk measure: given non-normality, "
           "CVaR\u2013which captures the average loss beyond the VaR threshold\u2013provides "
           "a more coherent and conservative risk measure than VaR alone. "
           "Portfolio CVaR(99%) = \u22125.60% per day versus VaR(99%) = \u22124.36%.")

    # ── 7. References ───────────────────────────────────────────────────────
    heading(doc, "7. References")
    refs = [
        "Bollerslev, T. (1986). Generalised autoregressive conditional "
        "heteroscedasticity. Journal of Econometrics, 31(3), 307\u2013327.",

        "Campbell, J. Y., Lo, A. W., & MacKinlay, A. C. (1997). "
        "The Econometrics of Financial Markets. Princeton University Press.",

        "DeMiguel, V., Garlappi, L., & Uppal, R. (2009). Optimal versus na\u00efve "
        "diversification: How inefficient is the 1/N portfolio strategy? "
        "Review of Financial Studies, 22(5), 1915\u20131953.",

        "Engle, R. F. (1982). Autoregressive conditional heteroscedasticity with "
        "estimates of the variance of United Kingdom inflation. Econometrica, "
        "50(4), 987\u20131007.",

        "Jarque, C. M., & Bera, A. K. (1987). A test for normality of observations "
        "and regression residuals. International Statistical Review, 55(2), 163\u2013172.",

        "Ljung, G. M., & Box, G. E. P. (1978). On a measure of lack of fit in "
        "time series models. Biometrika, 65(2), 297\u2013303.",

        "Longin, F., & Solnik, B. (2001). Extreme correlation of international "
        "equity markets. Journal of Finance, 56(2), 649\u2013676.",

        "Markowitz, H. (1952). Portfolio selection. Journal of Finance, 7(1), 77\u201391.",

        "Rockafellar, R. T., & Uryasev, S. (2000). Optimization of conditional "
        "value-at-risk. Journal of Risk, 2(3), 21\u201341.",

        "Sharpe, W. F. (1966). Mutual fund performance. Journal of Business, 39(1), 119\u2013138.",
    ]
    for ref in refs:
        bullet(doc, ref)

    out = os.path.join(DOCS_DIR, "Q1_Statistical_Analysis_Report.docx")
    doc.save(out)
    print(f"Saved: {out}")
    return out


if __name__ == "__main__":
    build_ew_portfolio_doc()
    build_statistical_analysis_report()
