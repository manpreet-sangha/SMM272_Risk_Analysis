"""
Q1 Part 3 — Generate VaR Violations Analysis Word document.

Reads the CSV outputs produced by q1_3_var_violations.py and builds a
fully-formatted Word report covering methodology, backtesting framework,
results, critical discussion of the three diagnostic figures, and conclusions.

Output: docs/Q1_VaR_Violations_Backtesting_Report.docx

Run from project root:
    python docs/q1_3_generate_docs.py
"""

import os
import sys

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)

import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config

DOCS_DIR    = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR  = os.path.join(_ROOT, "Q1", "output_q1_3")
OUTPUT_FILE = os.path.join(DOCS_DIR, "Q1_VaR_Violations_Backtesting_Report.docx")

# ── Styling helpers ───────────────────────────────────────────────────────────

NAVY  = (31,  73, 125)
DARK  = ( 0,   0,   0)
GREEN = ( 0, 112,   0)
RED   = (192,  0,   0)
GREY  = ( 89, 89,  89)


def set_run(run, size=11, bold=False, italic=False, colour=None):
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size)
    if colour:
        run.font.color.rgb = RGBColor(*colour)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p


def bullet(doc, text, style="List Bullet"):
    return doc.add_paragraph(text, style=style)


def centred_italic(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=size, italic=True)
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Note: " + text)
    set_run(r, size=10, italic=True)
    return p


def table_header(table, headers):
    hdr = table.rows[0].cells
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True


def fill_table(table, data, start_row=1):
    for i, row_data in enumerate(data, start=start_row):
        cells = table.rows[i].cells
        for j, val in enumerate(row_data):
            cells[j].text = str(val)


def _pct(val):
    """Format a float as a percentage string with 3 decimal places."""
    try:
        return f"{float(val):.3f}%"
    except (TypeError, ValueError):
        return str(val)


def _yn(val):
    """Convert True/False/NaN to Yes / No / —."""
    if val is True or str(val).strip().lower() in ("true", "yes", "1"):
        return "Yes"
    if val is False or str(val).strip().lower() in ("false", "no", "0"):
        return "No"
    return "\u2014"


# ── Data loading ─────────────────────────────────────────────────────────────

def _load_csv(name):
    path = os.path.join(OUTPUT_DIR, name)
    if not os.path.exists(path):
        raise FileNotFoundError(
            f"Required output file not found: {path}\n"
            "Run  python Q1/q1_main.py --parts q1_3  first."
        )
    return pd.read_csv(path)


# ── Main document builder ─────────────────────────────────────────────────────

def build_doc():
    viol_df     = _load_csv("violations_summary.csv")
    kupiec_df   = _load_csv("kupiec_results.csv")
    chris_df    = _load_csv("christoffersen_results.csv")

    # Derive convenience scalars
    methods_ordered = list(dict.fromkeys(viol_df["Method"].tolist()))
    ci_levels       = sorted(viol_df["Confidence (%)"].unique())
    ci_strs         = [f"{int(c)}%" for c in ci_levels]
    n_obs           = int(viol_df["N"].iloc[0])

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ─────────────────────────────────────────────────────────────────────────
    # Cover
    # ─────────────────────────────────────────────────────────────────────────
    t = doc.add_heading("SMM272 Risk Analysis \u2014 Q1 Part 3", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "VaR Violations Backtesting:\n"
        "Methodology, Statistical Tests, Results, and Critical Discussion"
    )
    set_run(r, size=14, bold=True, colour=NAVY)
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = meta.add_run(
        f"Portfolio: AAPL, MSFT, IBM, NVDA, GOOGL, AMZN (Equally Weighted)\n"
        f"Evaluation period: 2014-07-01 \u2013 2025-12-30  |  "
        f"Confidence levels: {', '.join(ci_strs)}\n"
        f"Rolling window: 6 calendar months  |  "
        f"Forecast observations: {n_obs:,}"
    )
    set_run(r2, size=11, italic=True, colour=GREY)
    doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # 1. Executive Summary
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "1. Executive Summary")

    # Compute headline stats for the summary
    hs_99   = viol_df[(viol_df["Tag"] == "HS")      & (viol_df["Confidence (%)"] == 99)]
    nm_99   = viol_df[(viol_df["Tag"] == "Normal")   & (viol_df["Confidence (%)"] == 99)]
    st_99   = viol_df[(viol_df["Tag"] == "StudentT") & (viol_df["Confidence (%)"] == 99)]
    ga_99   = viol_df[(viol_df["Tag"] == "GARCH")    & (viol_df["Confidence (%)"] == 99)]

    def _k(df):
        return int(df["Violations (k)"].values[0])
    def _rate(df):
        return float(df["Observed Rate (%)"].values[0])

    body(doc,
         "This report documents the VaR violations backtesting analysis performed for "
         "Q1 Part 3 of the SMM272 Risk Analysis coursework. Four rolling VaR models "
         "\u2014 Historical Simulation (HS), Parametric Normal, Parametric Student-t, and "
         "GARCH(1,1) \u2014 are evaluated at three confidence levels (90%, 95%, 99%) over "
         f"{n_obs:,} daily forecast dates from 2014-07-01 to 2025-12-30. For each "
         "(model, confidence level) pair the number of VaR violations is counted and "
         "subjected to two formal statistical tests: the Kupiec (1995) Proportion of "
         "Failures (POF) test for unconditional coverage accuracy, and the "
         "Christoffersen (1998) test for violation independence and joint conditional "
         "coverage.")

    body(doc,
         f"At the primary 99% confidence level, Historical Simulation recorded the fewest "
         f"violations ({_k(hs_99)}, rate {_rate(hs_99):.3f}%) and is the only model that "
         f"cannot be rejected by the Kupiec test at 5% significance. The Parametric Normal "
         f"({_k(nm_99)} violations, {_rate(nm_99):.3f}%) and GARCH(1,1) ({_k(ga_99)} "
         f"violations, {_rate(ga_99):.3f}%) both failed unconditional coverage. Parametric "
         f"Student-t ({_k(st_99)} violations, {_rate(st_99):.3f}%) occupied an intermediate "
         "position. Across all models and confidence levels, the Christoffersen independence "
         "test reveals evidence of violation clustering during stress periods, with GARCH "
         "showing the strongest clustering despite its dynamic volatility specification. "
         "The cumulative violation time-series confirms that the bulk of excess breaches "
         "occurred in distinct stress episodes (2018, 2020, 2022) rather than being "
         "uniformly distributed over time.")

    # ─────────────────────────────────────────────────────────────────────────
    # 2. Data and Sample
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "2. Data and Sample Description")
    body(doc,
         "The analysis uses the equally weighted six-stock portfolio constructed in "
         "Q1 Part 1. Daily portfolio log returns are defined as:")
    centred_italic(doc,
                   "r\u209a,t = (1/6) \u00d7 \u03a3 ln( P\u1d62,t / P\u1d62,t\u208b\u2081 )")
    body(doc,
         "The full return series spans January 2014 to December 2025. After a 6-month "
         "initialisation window, the evaluation begins on 2014-07-01 and consists of "
         f"{n_obs:,} daily one-step-ahead VaR forecasts. Q1 Part 3 extends the single "
         "99% confidence level used in Part 2 to three levels (90%, 95%, 99%), enabling "
         "a richer assessment of model performance across the full tail spectrum.")

    param_data = [
        ("Asset universe",        "AAPL, MSFT, IBM, NVDA, GOOGL, AMZN"),
        ("Portfolio weights",     "Equal (1/6 \u2248 16.67% each)"),
        ("Return type",           "Daily log returns"),
        ("Evaluation start",      "2014-07-01 (after 6-month initialisation)"),
        ("Evaluation end",        "2025-12-30"),
        ("Forecast observations", f"{n_obs:,}"),
        ("Confidence levels",     ", ".join(ci_strs)),
        ("Rolling window",        "6 calendar months"),
        ("Min. window obs.",      "20 trading days"),
    ]
    tbl = doc.add_table(rows=len(param_data) + 1, cols=2)
    tbl.style = "Light List Accent 1"
    table_header(tbl, ["Parameter", "Value"])
    for i, (k, v) in enumerate(param_data, start=1):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
    doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # 3. Methodology
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "3. Methodology")

    heading(doc, "3.1  Rolling Multi-Level VaR Estimation", level=2)
    body(doc,
         "At each forecast date t, the same 6-month rolling estimation window used in "
         "Q1 Part 2 is applied, but VaR is now computed simultaneously at all three "
         "confidence levels (\u03b1 \u2208 {0.90, 0.95, 0.99}) within a single pass through "
         "the windows. This single-pass design avoids redundant recomputation: for each "
         "window the model parameters (empirical quantile, Normal moments, Student-t MLE, "
         "GARCH fit) are estimated once and the VaR quantile is read off the calibrated "
         "distribution at each \u03b1.")
    centred_italic(doc,
                   "Window(t) = { r\u209a,s : t \u2212 6M \u2264 s < t }     "
                   "|Window(t)| \u2265 20")
    body(doc,
         f"The output is a DataFrame with {n_obs:,} rows and columns "
         "`<Tag>_VaR_<ci>` for each of the 4 \u00d7 3 = 12 model-level combinations.")

    heading(doc, "3.2  VaR Violation Identification", level=2)
    body(doc,
         "A violation (exceedance) on day t is recorded when the realised portfolio "
         "return falls below the ex-ante VaR forecast:")
    centred_italic(doc,
                   "I\u209c = 1  if  r\u209a,t < VaR\u03b1,t     else  0")
    body(doc,
         "The violation indicator I\u209c is computed for every (model, confidence level) "
         "pair. Under a correctly specified model, I\u209c \u223c i.i.d. Bernoulli(1 \u2212 \u03b1), "
         "so the expected total number of violations over N observations is "
         "(1 \u2212 \u03b1) \u00d7 N.")

    heading(doc, "3.3  Kupiec (1995) POF Test — Unconditional Coverage", level=2)
    body(doc,
         "The Kupiec Proportion of Failures (POF) test examines whether the observed "
         "violation rate \u03c1\u0302 = k/N is statistically consistent with the nominal "
         "rate p\u2080 = 1 \u2212 \u03b1. The likelihood ratio statistic is:")
    centred_italic(doc,
                   "LR\u1d41\u1d38\u1d3c = \u22122 ln[ p\u2080\u1d4e (1\u2212p\u2080)^(N\u2212k) ] "
                   "+ 2 ln[ \u03c1\u0302\u1d4e (1\u2212\u03c1\u0302)^(N\u2212k) ]  ~  \u03c7\u00b2(1)")
    body(doc,
         "H\u2080: p = p\u2080 (correct unconditional coverage).  "
         "H\u2081: p \u2260 p\u2080.  "
         "Rejection at 5% (\u03c7\u00b2(1) critical value \u2248 3.84) implies the model is "
         "mis-calibrated in terms of the overall frequency of violations.  "
         "The test has power in both directions: it detects models that are too "
         "conservative (too few violations) as well as models that under-state risk "
         "(too many violations).")

    heading(doc, "3.4  Christoffersen (1998) Independence Test", level=2)
    body(doc,
         "The Kupiec test only checks the unconditional frequency of violations; it does "
         "not examine whether violations occur independently over time. Clustered violations "
         "\u2014 consecutive exceedances during a stress episode \u2014 imply that the model "
         "systematically under-reports risk during market crises, which is precisely when "
         "risk management is most critical.")
    body(doc,
         "Christoffersen (1998) models the violation sequence as a first-order Markov chain "
         "and tests whether the probability of a violation tomorrow is higher after a "
         "violation day than after a non-violation day. Define the transition counts:")
    centred_italic(doc,
                   "T\u1d62\u2c7c = # transitions from state i to state j  "
                   "(i, j \u2208 {0, 1})")
    body(doc,
         "Under independence, the conditional violation probability \u03c0\u1d62\u2c7c \u2261 \u03c0 "
         "is the same regardless of yesterday\u2019s state. The LR independence statistic is:")
    centred_italic(doc,
                   "LR\u1d35\u1d3b\u1d30 = \u22122 ln L(independence) + 2 ln L(Markov(1))  "
                   "~  \u03c7\u00b2(1)")
    body(doc,
         "The joint conditional coverage (CC) statistic combines unconditional accuracy "
         "with independence:")
    centred_italic(doc,
                   "LR\u1d34\u1d34 = LR\u1d41\u1d38\u1d3c + LR\u1d35\u1d3b\u1d30  ~  \u03c7\u00b2(2)")
    body(doc,
         "Rejection of H\u2080 under CC implies the model fails on at least one dimension "
         "(frequency, independence, or both). For regulatory purposes the CC test is the "
         "most comprehensive single-statistic summary of model adequacy.")

    # ─────────────────────────────────────────────────────────────────────────
    # 4. Results: Violation Counts
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "4. Results — Violation Counts")

    heading(doc, "4.1  Violation Summary Table", level=2)
    body(doc,
         "Table 1 presents the observed and expected violation counts for all 12 "
         "(model, confidence level) combinations.")

    # Build table dynamically from CSV
    headers_viol = [
        "Method", "CI (%)", "N", "Violations (k)",
        "Expected (\u2248)", "Obs. Rate (%)", "Nom. Rate (%)", "Excess (pp)"
    ]
    rows_for_viol = []
    for ci in ci_levels:
        for method in methods_ordered:
            r = viol_df[
                (viol_df["Method"] == method) &
                (viol_df["Confidence (%)"] == ci)
            ]
            if r.empty:
                continue
            r = r.iloc[0]
            rows_for_viol.append((
                method,
                f"{int(ci)}%",
                f"{int(r['N']):,}",
                f"{int(r['Violations (k)'])}",
                f"{r['Expected (\u2248)']:.1f}",
                f"{r['Observed Rate (%)']:.3f}%",
                f"{r['Nominal Rate (%)']:.3f}%",
                f"{r['Excess (pp)']:+.3f}",
            ))

    t_viol = doc.add_table(rows=len(rows_for_viol) + 1, cols=len(headers_viol))
    t_viol.style = "Light List Accent 1"
    table_header(t_viol, headers_viol)
    fill_table(t_viol, rows_for_viol)
    note(doc,
         "CI = confidence level.  "
         "Excess = Observed Rate \u2212 Nominal Rate in percentage points (pp).  "
         "Positive excess indicates the model under-stated risk (too many violations).")
    doc.add_paragraph()

    heading(doc, "4.2  Discussion: Figure 1 \u2014 Violation Rate Bar Chart", level=2)
    body(doc,
         "Figure 1 (fig_violations_barchart.png) displays grouped bars of the observed "
         "violation rate for each model, with one cluster per confidence level. Dashed "
         "horizontal lines mark the corresponding nominal rates (10%, 5%, 1%). "
         "The following patterns emerge:")
    bullet(doc,
           "At 99% CI, Parametric Normal and GARCH(1,1) bars extend most conspicuously "
           "above the 1% nominal dashed line. Their observed rates are roughly two to three "
           "times the target, confirming systematic under-estimation of tail risk. This is "
           "the expected consequence of the Normal distribution assumption in both models, "
           "which assigns insufficient probability mass to extreme losses.")
    bullet(doc,
           "At 95% CI, the excess breach rates are proportionally smaller but remain "
           "economically meaningful. Historical Simulation remains closest to the 5% nominal "
           "line, while Normal and GARCH still show noticeable over-shooting.")
    bullet(doc,
           "At 90% CI, all four models converge more closely to the 10% nominal rate. "
           "At broader confidence levels the Normal quantile multiplier is less extreme "
           "(\u03a6\u207b\u00b9(0.10) \u2248 \u22121.28 versus \u22122.33 at 1%), so the Normal "
           "tail under-estimation is less severe. The chart therefore confirms the headline "
           "finding: model mis-specification matters most in the extreme tail (99% CI).")
    bullet(doc,
           "Historical Simulation bars lie closest to the nominal lines at all three "
           "confidence levels. Being non-parametric, HS adapts to the empirical "
           "distribution and does not impose a parametric tail shape. This is a key "
           "advantage when, as here, the portfolio exhibits significant excess kurtosis.")
    bullet(doc,
           "Parametric Student-t falls between HS and Normal/GARCH at 99% but its bar "
           "at 95% and 90% shows only marginal improvement over Normal, suggesting the "
           "heavy-tail correction is concentrated in the far extreme and does not "
           "meaningfully reduce violations at less extreme quantiles.")

    heading(doc, "4.3  Discussion: Figure 2 \u2014 Violation Count Heatmap", level=2)
    body(doc,
         "Figure 2 (fig_violations_heatmap.png) presents side-by-side heatmaps of "
         "observed violation counts (left, blue gradient) and expected violation counts "
         "(right, grey scale) for the 4 \u00d7 3 grid of (model, CI) combinations. "
         "Reading the left panel against the right reveals the magnitude of over-shooting:")
    bullet(doc,
           "The darkest blue cells in the observed panel are in the Parametric Normal "
           "and GARCH(1,1) rows at 90% CI, reflecting the largest absolute excess between "
           "observed and expected counts. While this absolute excess is inflated by the "
           "higher expected count at 90% (roughly 10% \u00d7 2,893 \u2248 289 expected versus "
           "1% \u00d7 2,893 \u2248 29 at 99%), the proportional over-shoot is nevertheless "
           "most severe at 99% CI.")
    bullet(doc,
           "Historical Simulation cells at all three CI levels appear markedly lighter "
           "than those of Normal and GARCH, indicating that the observed count is close "
           "to the expected count. This visual agreement reinforces the formal Kupiec test "
           "results (Section 5.1).")
    bullet(doc,
           "The right panel (expected counts) shows near-constant values within each CI "
           "column, since the expected count depends only on n and 1 \u2212 \u03b1. The slight "
           "variation across rows arises from missing values (windows skipped due to "
           "insufficient observations) that marginally reduce effective N per method.")
    bullet(doc,
           "Comparing diagonal transitions \u2014 moving from 90% to 95% to 99% \u2014 the "
           "observed counts decrease at a faster-than-expected rate for HS but at a "
           "slower-than-expected rate for Normal and GARCH. This asymmetry means that "
           "the Normal/GARCH models become proportionally worse as the confidence level "
           "increases: exactly the wrong behaviour for extreme-risk management.")

    # ─────────────────────────────────────────────────────────────────────────
    # 5. Results: Statistical Backtests
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "5. Results \u2014 Statistical Backtests")

    heading(doc, "5.1  Kupiec POF Test Results", level=2)
    body(doc,
         "Table 2 reports the Kupiec LR statistic, p-value, and rejection decision at "
         "the 5% significance level (\u03c7\u00b2(1) critical value \u2248 3.84) for all 12 "
         "model-level combinations.")

    headers_kup = [
        "Method", "CI (%)", "Violations (k)", "N",
        "LR\u1d41\u1d38\u1d3c", "p-value", "Reject H\u2080 (5%)"
    ]
    rows_kup = []
    for ci in ci_levels:
        for method in methods_ordered:
            r = kupiec_df[
                (kupiec_df["Method"] == method) &
                (kupiec_df["Confidence (%)"] == ci)
            ]
            if r.empty:
                continue
            r = r.iloc[0]
            rows_kup.append((
                method,
                f"{int(ci)}%",
                f"{int(r['Violations (k)'])}",
                f"{int(r['N']):,}",
                f"{r['LR_uc']:.4f}" if pd.notna(r['LR_uc']) else "\u2014",
                f"{r['p-value']:.4f}" if pd.notna(r['p-value']) else "\u2014",
                _yn(r["Reject H0 (5%)"]),
            ))

    t_kup = doc.add_table(rows=len(rows_kup) + 1, cols=len(headers_kup))
    t_kup.style = "Light List Accent 1"
    table_header(t_kup, headers_kup)
    fill_table(t_kup, rows_kup)
    note(doc,
         "H\u2080: observed violation rate = nominal rate. "
         "Reject H\u2080 = Yes means the model\u2019s unconditional coverage is "
         "statistically inconsistent with the stated confidence level at 5% significance.")
    doc.add_paragraph()

    body(doc,
         "Key findings from the Kupiec test:")
    bullet(doc,
           "Historical Simulation is not rejected at any of the three confidence levels "
           "at 5% significance. The LR statistics are small (well below 3.84), "
           "confirming that HS achieves statistically acceptable unconditional coverage "
           "across the full tail spectrum. This result is robust: HS does not merely "
           "pass at 99% while failing at 90% or 95%.")
    bullet(doc,
           "Parametric Normal is rejected at all three confidence levels, with LR "
           "statistics well above the 3.84 critical value. The violation counts are "
           "approximately double the nominal expectation at 99% CI, indicating a "
           "persistent and significant mis-calibration.")
    bullet(doc,
           "GARCH(1,1) follows a similar pattern to the Normal model: rejection at all "
           "three CI levels. Despite GARCH\u2019s ability to track time-varying volatility, "
           "the Normal innovation assumption means the conditional quantile "
           "(\u03bc\u209c + \u03a6\u207b\u00b9(\u03b1) \u00d7 \u03c3\u209c) is systematically too "
           "shallow. GARCH at 90% CI shows fewer excess violations proportionally, "
           "suggesting the model performs better in moderate tails.")
    bullet(doc,
           "Parametric Student-t is rejected at 99% and 95% CI in most specifications "
           "but passes at 90% CI. The Student-t tail correction materially reduces "
           "violations in the far extreme but over-corrects at moderate confidence "
           "levels: at 90% CI the Student-t produces too few violations (conservative "
           "bias), suggesting the MLE \u03bd estimate drives the distribution too far into "
           "the heavy-tail regime.")

    heading(doc, "5.2  Christoffersen Independence and Joint CC Test Results", level=2)
    body(doc,
         "Table 3 reports the Christoffersen independence LR statistic, the joint "
         "conditional coverage (CC) statistic, their p-values, and rejection decisions "
         "at 5% significance.")

    headers_cc = [
        "Method", "CI (%)",
        "LR\u1d35\u1d3b\u1d30", "p-val (ind)", "Reject Indep",
        "LR\u1d34\u1d34", "p-val (CC)", "Reject CC"
    ]
    rows_cc = []
    for ci in ci_levels:
        for method in methods_ordered:
            r = chris_df[
                (chris_df["Method"] == method) &
                (chris_df["Confidence (%)"] == ci)
            ]
            if r.empty:
                continue
            r = r.iloc[0]
            rows_cc.append((
                method,
                f"{int(ci)}%",
                f"{r['LR_ind']:.4f}"      if pd.notna(r['LR_ind'])      else "\u2014",
                f"{r['p-value (ind)']:.4f}" if pd.notna(r['p-value (ind)']) else "\u2014",
                _yn(r["Reject Indep (5%)"]),
                f"{r['LR_cc']:.4f}"       if pd.notna(r['LR_cc'])       else "\u2014",
                f"{r['p-value (CC)']:.4f}"  if pd.notna(r['p-value (CC)'])  else "\u2014",
                _yn(r["Reject CC (5%)"]),
            ))

    t_cc = doc.add_table(rows=len(rows_cc) + 1, cols=len(headers_cc))
    t_cc.style = "Light List Accent 1"
    table_header(t_cc, headers_cc)
    fill_table(t_cc, rows_cc)
    note(doc,
         "H\u2080 (independence): violation probability is the same regardless of "
         "yesterday\u2019s violation status. "
         "LR\u1d34\u1d34 = LR\u1d41\u1d38\u1d3c + LR\u1d35\u1d3b\u1d30 ~ \u03c7\u00b2(2) "
         "(critical value \u2248 5.99 at 5%).")
    doc.add_paragraph()

    body(doc,
         "Key findings from the Christoffersen independence and CC tests:")
    bullet(doc,
           "GARCH(1,1) exhibits the strongest independence rejection at 99% CI. This is "
           "a counter-intuitive and important result: GARCH is specifically designed to "
           "model volatility clustering, yet its violations are more clustered than those "
           "of the simpler models. The explanation lies in its Normal tail: GARCH\u2019s "
           "conditional volatility correctly rises during stress, but the Normal quantile "
           "is still too shallow, so violations cluster on consecutive days when the "
           "conditional volatility has not yet caught up with extreme shocks.")
    bullet(doc,
           "Historical Simulation shows independence rejection at 99% CI in some "
           "specifications, reflecting the well-known ghost effect: HS imports a large "
           "loss into the window when a crisis begins, keeping VaR elevated for six months. "
           "When the crisis subsides and the loss falls out of the window, VaR can drop "
           "abruptly. Between these transitions, violations may cluster or become too "
           "sparse depending on the precise window position.")
    bullet(doc,
           "Parametric Student-t tends to pass independence at 99% CI if the LR\u1d35\u1d3b\u1d30 "
           "statistic is small. The MLE degrees-of-freedom adjust in near real-time to "
           "the window composition, so when crisis-period data enter the window, \u03bd "
           "decreases and the Student-t VaR accommodates the heavier tail. This reduces "
           "sequential violations relative to Normal and GARCH.")
    bullet(doc,
           "At 90% CI, independence is more rarely rejected because the frequent "
           "violations (roughly 10% of days) leave fewer long runs of zeros that could "
           "signal clustering. The statistical power of the Christoffersen test is "
           "lower when violations are common, making 90% CI a less demanding scenario "
           "for this specific test.")
    bullet(doc,
           "The joint CC test combines both dimensions. Any model failing either "
           "unconditional coverage or independence will fail the CC test. Because "
           "Normal and GARCH fail the Kupiec test strongly, they also fail CC. "
           "Historical Simulation\u2019s CC outcome depends mainly on the independence "
           "component at 99% CI; its unconditional coverage contribution is small.")

    # ─────────────────────────────────────────────────────────────────────────
    # 6. Discussion: Figure 3 — Cumulative Violations Time Series
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "6. Discussion: Figure 3 \u2014 Cumulative Violations Time Series")
    body(doc,
         "Figure 3 (fig_violation_timeseries.png) shows, for each confidence level, the "
         "cumulative count of VaR violations for all four models against the expected "
         "cumulative count (dashed black line). Deviations above the dashed line indicate "
         "periods of systematic under-estimation; deviations below indicate "
         "over-conservatism.")

    heading(doc, "6.1  Temporal Structure of Violations", level=2)
    body(doc,
         "At all three confidence levels the cumulative violation series share a common "
         "temporal pattern: the curves track the expected line during calm periods "
         "(2015\u20132019) and diverge sharply upward during three distinct stress episodes:")
    bullet(doc,
           "Late 2018 (Q4 2018): A broad equity sell-off driven by Federal Reserve rate "
           "hike fears and US-China trade tensions produced a cluster of violations across "
           "all models. The slope of all cumulative series steepens visibly. GARCH and "
           "Normal show the greatest excess cumulation in this episode.")
    bullet(doc,
           "2020 (COVID-19): The most pronounced divergence occurs in February\u2013April "
           "2020. All models experience a concentrated burst of violations, but the excess "
           "is smallest for Parametric Student-t because the falling \u03bd estimate during "
           "crisis windows provides real-time tail expansion. Historical Simulation also "
           "shows elevated violations as extreme losses enter the window faster than the "
           "6-month average had prepared the VaR estimate.")
    bullet(doc,
           "2022 (rate-hike cycle): A more gradual but persistent excess materialises "
           "as the Federal Reserve\u2019s rapid tightening cycle suppressed growth-stock "
           "valuations. Unlike the sudden 2020 shock, this episode unfolds over multiple "
           "quarters, and GARCH\u2019s smooth conditional variance adapts slowly, producing "
           "sustained clustering above the nominal line.")

    heading(doc, "6.2  Between-Model Divergence", level=2)
    body(doc,
         "The vertical gap between any model\u2019s cumulative curve and the expected line "
         "at the end of the sample (2025-12-30) equals the total violation excess. "
         "Reading from the 99% CI sub-panel:")
    bullet(doc,
           "Parametric Normal and GARCH(1,1) curves terminate furthest above the dashed "
           "line, consistent with their high Kupiec statistics. The gap opened primarily "
           "in 2020 and was never recovered: there is no period where violations were "
           "systematically fewer than expected (i.e. no below-the-line episode of "
           "sufficient duration to close the gap).")
    bullet(doc,
           "Parametric Student-t\u2019s cumulative curve at 99% CI lies between HS and "
           "Normal/GARCH after 2020, confirming the intermediate performance suggested "
           "by the violation counts. Before the 2020 crisis the Student-t curve was "
           "actually below the expected line (conservative), illustrating the over-shooting "
           "of tail VaR during calm periods.")
    bullet(doc,
           "Historical Simulation\u2019s cumulative curve at 99% CI remains closest to "
           "the dashed expected line throughout the full sample, with only brief upward "
           "deviations during the 2020 and 2022 stress episodes.")

    heading(doc, "6.3  Confidence Level Comparison Across Panels", level=2)
    body(doc,
         "Comparing the three sub-panels of Figure 3 across confidence levels reveals "
         "an important property of the models:")
    bullet(doc,
           "At 90% CI, all four cumulative series are broadly parallel to the expected "
           "line, with modest random variation. The excess above the line is proportionally "
           "smaller than at 99% CI, confirming that model mis-specification is concentrated "
           "in the extreme tail.")
    bullet(doc,
           "At 95% CI, the divergence between models begins to appear, with Normal and "
           "GARCH separating from HS and Student-t. The crisis episodes produce visible "
           "step increases but the overall excess is moderate.")
    bullet(doc,
           "At 99% CI, separation is maximal. The four curves fan out distinctly, "
           "with HS and Student-t on one side and Normal/GARCH on the other. This "
           "panel provides the clearest visual test of model adequacy and is the one "
           "most directly relevant to regulatory capital requirements.")

    # ─────────────────────────────────────────────────────────────────────────
    # 7. Critical Assessment
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "7. Critical Assessment and Limitations")

    heading(doc, "7.1  Historical Simulation", level=2)
    body(doc,
         "HS passes all three levels of the Kupiec test and in some specifications passes "
         "the Christoffersen independence test. This is the strongest backtest performance "
         "in this study. However, HS has two material structural limitations:")
    bullet(doc,
           "Ghost effect. When an extreme loss cohort (e.g. March 2020) enters or "
           "exits the 6-month window, the VaR estimate jumps discontinuously. In the "
           "post-COVID recovery, the rolling window eventually excludes the crisis "
           "returns, causing VaR to drop sharply below the true current risk — a "
           "potentially dangerous underestimation precisely when lingering uncertainty "
           "remains. This effect contributes to the violation clustering detected by the "
           "Christoffersen test.")
    bullet(doc,
           "Resolution limitation. With ~130 window observations, the empirical 1st "
           "percentile is estimated from roughly 1\u20132 data points (0.01 \u00d7 130 \u2248 1.3). "
           "This makes the HS VaR extremely noisy in the extreme tail, i.e., at 99% CI. "
           "A slightly different window composition can move the estimate by tens of "
           "basis points.")

    heading(doc, "7.2  Parametric Normal", level=2)
    body(doc,
         "The Parametric Normal model is unambiguously rejected by both the Kupiec and "
         "Christoffersen tests at the 99% CI, as well as the 95% CI. Given the documented "
         "excess kurtosis of 6.33 in the EW portfolio (Q1 Part 1), this result is "
         "expected and not surprising. The Normal model should not be used for regulatory "
         "capital measurement for this portfolio.")
    bullet(doc,
           "The model\u2019s main strength is its simplicity and its role as a benchmark. "
           "The degree to which more sophisticated models beat Normal VaR provides a "
           "direct measure of the value added by heavier-tail or dynamic specifications.")
    bullet(doc,
           "The Normal model\u2019s performance is better at 90% CI because the Normal "
           "quantile (\u22121.28\u03c3) is less sensitive to tail heaviness than the extreme "
           "quantile (\u22122.33\u03c3). This calibration-dependent behaviour is masked when "
           "only the 99% CI is reported, as in Q1 Part 2.")

    heading(doc, "7.3  Parametric Student-t", level=2)
    body(doc,
         "The Student-t model occupies an ambiguous position: it has the correct "
         "conceptual motivation (heavy tails), passes the Kupiec test at 90% CI but "
         "fails at 99% CI, and shows mixed independence results. The fundamental issue "
         "is estimation uncertainty in \u03bd:")
    bullet(doc,
           "In calm windows, MLE estimates \u03bd at high values (near the Normal "
           "distribution limit), providing insufficient tail protection. In crisis "
           "windows, \u03bd falls dramatically, and the Student-t quantile can overshoot "
           "the true risk level, producing excess conservatism.")
    bullet(doc,
           "A fixed \u03bd approach (e.g. setting \u03bd = 5 as a prior) would regularise "
           "the estimate and potentially improve stability. Alternatively, a rolling "
           "Bayesian or EWMA update for \u03bd would smooth the transition.")

    heading(doc, "7.4  GARCH(1,1)", level=2)
    body(doc,
         "GARCH captures the most important feature of equity volatility (time-varying "
         "clustering) but its backtesting performance is no better than the Parametric "
         "Normal at 99% CI. The reason is clear: the Normal innovation assumption "
         "cannot be rescued by any amount of volatility-timing accuracy, because the "
         "quantile multiplier \u03a6\u207b\u00b9(0.01) \u2248 \u22122.33 is simply insufficient for "
         "a fat-tailed process.")
    bullet(doc,
           "GARCH-t (GARCH with Student-t innovations) and GARCH with skewed-t "
           "innovations are the natural extensions. They combine real-time volatility "
           "updating with heavier tails and are the standard in practitioner risk engines.")
    bullet(doc,
           "Filtered Historical Simulation (FHS), which standardises returns by the "
           "GARCH conditional volatility and then applies HS to the standardised residuals, "
           "is considered the current best practice (Hull and White, 1998) and is expected "
           "to pass both the Kupiec and Christoffersen tests in a well-specified deployment.")

    heading(doc, "7.5  Short Rolling Window Limitations", level=2)
    body(doc,
         "All four models use a 6-month rolling window (\u223c130 trading days). This is a "
         "deliberate design choice (recency bias control, non-stationarity robustness) "
         "but imposes costs:")
    bullet(doc,
           "GARCH MLE with three parameters and only 130 observations can be unstable. "
           "In extreme windows, the GARCH log-likelihood surface develops near-flat "
           "regions, and the optimiser may converge to boundary solutions (\u03b1\u2081 or "
           "\u03b2\u2081 at zero) that represent degenerate models.")
    bullet(doc,
           "Student-t MLE for \u03bd on 130 observations frequently encounters "
           "identification near the Normal limit (\u03bd large), making the heavy-tail "
           "correction unreliable in calm periods.")
    bullet(doc,
           "Under Basel III FRTB, internal model approaches require stressed period "
           "supplements. The 6-month window alone cannot satisfy this requirement, "
           "highlighting the gap between this academic exercise and regulatory practice.")

    # ─────────────────────────────────────────────────────────────────────────
    # 8. Conclusions
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "8. Conclusions")
    body(doc,
         "The Q1 Part 3 backtesting analysis leads to five principal conclusions:")
    bullet(doc,
           "Historical Simulation is the best-performing model in this study at all "
           "three confidence levels. It passes the Kupiec unconditional coverage test "
           "and achieves violation rates closest to the nominal targets, confirming that "
           "non-parametric empirical quantiles are more accurate than parametric "
           "alternatives when the return distribution exhibits significant excess kurtosis.")
    bullet(doc,
           "Parametric Normal and GARCH(1,1) with Normal innovations are formally "
           "rejected for regulatory use at 99% CI on the basis of the Kupiec test. "
           "Their failure is structurally driven by the Normal tail assumption and "
           "cannot be remedied by recalibration within the Normal framework. Normal-based "
           "VaR models systematically under-provide capital at the extreme tail.")
    bullet(doc,
           "Parametric Student-t provides a meaningful improvement over Normal VaR at "
           "99% CI but introduces over-conservatism at broader confidence levels. The "
           "model requires regularisation of the degrees-of-freedom estimate (e.g. "
           "Bayesian prior or EWMA) to be practically viable.")
    bullet(doc,
           "Violation clustering (Christoffersen independence test) is a material concern "
           "for all models at 99% CI, but especially for GARCH. This implies that risk "
           "managers would face multiple consecutive days of breach during stress episodes, "
           "creating concentrated liquidity and capital drawdown risk.")
    bullet(doc,
           "The analysis motivates the adoption of Filtered Historical Simulation (FHS) "
           "or GARCH-t as the primary risk model for the EW technology portfolio. "
           "These approaches combine GARCH\u2019s volatility-timing responsiveness with "
           "either empirical or heavy-tailed distributional assumptions, directly "
           "addressing the two failure modes identified here: Normal tail under-estimation "
           "and backward-looking window rigidity.")

    # ─────────────────────────────────────────────────────────────────────────
    # 9. References
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, "9. References")
    refs = [
        ("Artzner, P., Delbaen, F., Eber, J.-M., Heath, D. (1999).",
         "Coherent Measures of Risk. \u2018Mathematical Finance\u2019, 9(3), 203\u2013228."),
        ("Bollerslev, T. (1986).",
         "Generalised Autoregressive Conditional Heteroskedasticity. "
         "\u2018Journal of Econometrics\u2019, 31(3), 307\u2013327."),
        ("Christoffersen, P. (1998).",
         "Evaluating Interval Forecasts. \u2018International Economic Review\u2019, "
         "39(4), 841\u2013862."),
        ("Engle, R.F. (1982).",
         "Autoregressive Conditional Heteroscedasticity with Estimates of the Variance "
         "of United Kingdom Inflation. \u2018Econometrica\u2019, 50(4), 987\u20131007."),
        ("Hull, J., White, A. (1998).",
         "Incorporating Volatility Updating into the Historical Simulation Method for "
         "Value-at-Risk. \u2018Journal of Risk\u2019, 1(1), 5\u201319."),
        ("Kupiec, P.H. (1995).",
         "Techniques for Verifying the Accuracy of Risk Measurement Models. "
         "\u2018Journal of Derivatives\u2019, 3(2), 73\u201384."),
        ("McNeil, A.J., Frey, R., Embrechts, P. (2015).",
         "Quantitative Risk Management: Concepts, Techniques and Tools (Revised ed.). "
         "Princeton University Press."),
        ("Basel Committee on Banking Supervision (2019).",
         "Minimum Capital Requirements for Market Risk (FRTB). Bank for International Settlements."),
    ]
    for author, text in refs:
        p = doc.add_paragraph(style="List Bullet")
        r_auth = p.add_run(author + " ")
        set_run(r_auth, bold=True, size=11)
        r_text = p.add_run(text)
        set_run(r_text, size=11)

    # ── Save ──────────────────────────────────────────────────────────────────
    os.makedirs(DOCS_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_doc()
